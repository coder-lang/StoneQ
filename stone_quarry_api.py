#stone_final
#output_generation changed to output_generation1 and port change to 8505.
#original working is output_generation.py and port is 8504
from fastapi import FastAPI, UploadFile, File, HTTPException, BackgroundTasks
from fastapi.responses import FileResponse, JSONResponse
from typing import Dict, Any, List, Tuple
import zipfile
import tempfile
import os
import uuid
import logging
import shutil
from pathlib import Path
from dotenv import load_dotenv
from fastapi.middleware.cors import CORSMiddleware 
from azure.storage.blob import BlobServiceClient, generate_blob_sas, BlobSasPermissions
import json
import re
from urllib.parse import unquote
from datetime import datetime, timedelta
import io
from functions.verification import verify_documents
from docx import Document

# ✅ NEW IMPORT
from file_filter import unzip_files_combined_filtered
from functions.seac_extraction import (
    unzip_seac_pdfs,
    extract_all_seac_docs,
)

# Load environment variables
load_dotenv()
# Configure logging
logging.basicConfig(level=os.getenv("LOG_LEVEL", "INFO"))
logger = logging.getLogger(__name__)
# Import your modules
from functions.doc_intelligence import extract_markdown_from_file
from functions.entity_extraction import (
    extract_caf, extract_form1, extract_nabet, extract_cluster_certificate,
    extract_site_survey, extract_emp, extract_mpa, extract_nocgp,
    extract_nocforest, extract_nocgsda, extract_kprat, extract_gsr,
    extract_qlp, extract_od, kml_to_json, safe_json_parse, extract_dsr, extract_regrassing, extract_undertaking, extract_western_ghat, extract_8A, extract_pfr, extract_form1A,extract_form1B,
    extract_unproponent, extract_unconsultant
)
from functions.output_generation12 import (
    gen_delib_sheet,
    fill_word_with_mapping,
    fill_mom_from_info_and_delib,
    build_info_sheet_mapping
)

from parallel_processor import (
    ParallelDocumentProcessor,
    build_file_mapping,
    build_verification_pairs
)

app = FastAPI(title="Document Verification Engine API", version="1.0.0")
app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "http://localhost:4200",
        "http://localhost:3000",
        "https://10.40.108.197",
        "http://10.40.108.197",
        "https://10.40.108.197:8504",
        "http://10.40.108.197:8504",
        "https://stonequarry.ey.net:8504",
        "https://stonequarry.ey.net", 
        "https://stonequarry.ey.net/ai/validate?",
        "*",
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
    allow_origin_regex=r"https?://.*\.?10\.40\.108\.197.*"
)

# Configuration
ZIP_UPLOAD_DIR = Path(os.getenv("ZIP_UPLOAD_DIR", "./uploads"))
OUTPUT_DIR = Path(os.getenv("OUTPUT_DIR", "./outputs"))
LOG_DIR = Path(os.getenv("LOG_DIR", "./logs"))
TEMPLATE_PATH = Path(os.getenv("TEMPLATE_PATH", "./templates/New Info Sheet Format_3.12.25_Final.docx"))

# Ensure directories exist
ZIP_UPLOAD_DIR.mkdir(exist_ok=True)
OUTPUT_DIR.mkdir(exist_ok=True)
LOG_DIR.mkdir(exist_ok=True)

# Azure Blob Storage Configuration
account_name = os.getenv("AZURE_STORAGE_ACCOUNT_NAME")
account_key = os.getenv("AZURE_STORAGE_ACCOUNT_KEY")
container_name = "aimhproposals"

# Initialize Azure Blob Service Client
blob_service_client = BlobServiceClient(
    account_url=f"https://{account_name}.blob.core.windows.net",
    credential=account_key
)

# Define the list of REQUIRED document types for validation
REQUIRED_DOC_NAMES = [
    "CAF", "Form 1", "NABET", "Cluster Certificate", "DMO Site Survey",
    "EMP", "MPA", "NOC-GP", "NOC-GSDA", "Kprat", "Quarry Layout Plan", "Ownership Document", "DSR", "Regrassing", "Undertaking", "Affidavit Proposal", "8A", "form1A", "form1B", "pfr", "un_consultant", "un_proponent"
]

# Global store for extraction results and blob URLs
extraction_results: Dict[str, Dict[str, Any]] = {}
blob_urls: Dict[str, Dict[str, str]] = {}
verification_results: Dict[str, Dict[Tuple[str, str], Any]] = {}
missing_files_store: Dict[str, List[str]] = {}

seac_results: Dict[str, Dict[str, Any]] = {}        # proposal_id → {filename: extracted_json}
seac_processing_status: Dict[str, str] = {}         # proposal_id → "Processing" | "Completed" | "Failed"

def safe_filename(proposal_id: str) -> str:
    """
    Convert any proposal_id into a filesystem-safe filename.
    Preserves original ID but returns a sanitized version for disk use.
    Example: 'SIA/MIN/4353/2024' → 'SIA_MIN_4353_2024'
    """
    return re.sub(r'[\\/:\*\?"<>\|\s]+', '_', proposal_id.strip())

# --- API Endpoints ---

@app.get("/")
def root():
    return {"message": "Document Verification Engine API is running"}

@app.get("/health")
def health_check():
    return {"status": "ok"}

@app.post("/validate")
async def validate_documents(
    proposal_id: str,
    zip_file: UploadFile = File(...)
):
    """
    OPTIMIZED: Check whether required documents are present in the ZIP.
    Now uses parallel file extraction for faster validation.
    """
    import time
    start_time = time.time()
    
    if not zip_file.filename.endswith('.zip'):
        raise HTTPException(status_code=400, detail="File must be a ZIP archive.")

    original_proposal_id = unquote(proposal_id)
    safe_id = safe_filename(original_proposal_id)
    zip_path = ZIP_UPLOAD_DIR / f"{safe_id}.zip"

    try:
        # Save ZIP file
        with open(zip_path, 'wb') as f:
            f.write(await zip_file.read())

        # Create proposal directory
        proposal_dir = OUTPUT_DIR / safe_id
        proposal_dir.mkdir(exist_ok=True)

        # Save metadata
        metadata = {
            "original_proposal_id": original_proposal_id,
            "upload_timestamp": str(datetime.now())
        }
        with open(proposal_dir / "metadata.json", "w") as meta_file:
            json.dump(metadata, meta_file, indent=2)

        # ✅ UPDATED: Use filtered extraction
        with tempfile.TemporaryDirectory() as temp_dir:
            logger.info(f"Extracting ZIP for validation: {safe_id}")
            
            # NEW: Filtered extraction (only relevant files)
            pdf_files, kml_files, ignored_count = unzip_files_combined_filtered(
                str(zip_path), 
                temp_dir
            )
            
            # Log if extra files were skipped
            if ignored_count > 0:
                logger.info(f"ℹ️  Skipped {ignored_count} non-matching files during validation")
            
            # Use build_file_mapping for consistent logic
            from parallel_processor import build_file_mapping
            file_mapping = build_file_mapping(pdf_files, kml_files)
            
            # Determine missing files
            found_files = {
                "CAF": file_mapping.get('caf'),
                "Form 1": file_mapping.get('form1'),
                "NABET": file_mapping.get('nabet'),
                "Cluster Certificate": file_mapping.get('cluster'),
                "DMO Site Survey": file_mapping.get('dmoss'),
                "EMP": file_mapping.get('emp'),
                "MPA": file_mapping.get('mpa'),
                "NOC-GP": file_mapping.get('nocgp'),
                "NOC-Forest": file_mapping.get('nocforest'),
                "NOC-GSDA": file_mapping.get('nocgsda'),
                "Kprat": file_mapping.get('kprat'),
                "GSR": file_mapping.get('gsr'),
                "Quarry Layout Plan": file_mapping.get('qlp'),
                "Ownership Document": file_mapping.get('od'),
                "DSR": file_mapping.get('dsr'),
                "Affidavit Proposal": file_mapping.get('western_ghat'),
                "8A": file_mapping.get('8A'),
                "form1A":file_mapping.get("form1A"),
                "form1B":file_mapping.get("form1B"),
                "pfr":file_mapping.get("pfr"),
                "un_consultant" : file_mapping.get("un_consultant"),
                "un_proponent" : file_mapping.get("un_proponent")
            }
            
            missing_docs = [
                name for name, file_path in found_files.items() 
                if name in REQUIRED_DOC_NAMES and file_path is None
            ]

        # Determine status
        if not missing_docs:
            status = "ALL Documents are Present"
            message = None
        else:
            status = "ALL Documents are not Present"
            message = ", ".join(missing_docs)

        elapsed = time.time() - start_time
        logger.info(f"✅ Validation completed in {elapsed:.2f}s for {original_proposal_id}")

        return {
            "status": status, 
            "message": message, 
            "proposal_id": original_proposal_id,
            "ignored_files_count": ignored_count  # NEW: Inform user about skipped files
        }
        
    except Exception as e:
        logger.error(f"Error during validation: {e}")
        if zip_path.exists():
            os.remove(zip_path)
        proposal_dir = OUTPUT_DIR / safe_id
        if proposal_dir.exists():
            shutil.rmtree(proposal_dir)
        raise HTTPException(status_code=500, detail=f"Validation failed: {str(e)}")

@app.post("/verify-process/{proposal_id:path}")
async def verify_process(proposal_id: str, background_tasks: BackgroundTasks = BackgroundTasks()):
    """
    This endpoint processes ALL detected documents using the ZIP file saved by /validate.
    It expects the proposal_id as a path parameter (e.g., /verify-process/SIA/MIN/4353/2024).
    """
    # The proposal_id received in path is already decoded by FastAPI when it contains slashes
    # So we can use it directly as the original proposal_id
    original_proposal_id = proposal_id
    # Use safe name for file access
    safe_id = safe_filename(original_proposal_id)

    # Construct the path to the saved ZIP file using safe_id
    zip_path = ZIP_UPLOAD_DIR / f"{safe_id}.zip"

    # Check if the ZIP file exists
    if not zip_path.exists():
        raise HTTPException(status_code=404, detail="Uploaded ZIP file not found for this proposal ID. Please run validation first.")

    # Create proposal directory to store metadata
    proposal_dir = OUTPUT_DIR / safe_id
    proposal_dir.mkdir(exist_ok=True)

    # Check for ongoing processing
    lock_file = proposal_dir / "processing.lock"
    if lock_file.exists():
        raise HTTPException(status_code=409, detail="Processing already in progress for this proposal ID.")

    # Create lock file
    lock_file.touch()

    log_dir = proposal_dir / "logs"
    log_dir.mkdir(exist_ok=True)

    extract_dir = proposal_dir / "extracted"
    extract_dir.mkdir(exist_ok=True)

    # Start background processing with the existing zip_path
    background_tasks.add_task(_process_zip_background, zip_path, original_proposal_id, proposal_dir, log_dir, lock_file)

    return {"proposal_id": original_proposal_id, "message": "Processing started in the background."}

def _process_zip_background(zip_path, original_proposal_id: str, proposal_dir, log_dir, lock_file):
    """
    OPTIMIZED background task with parallel processing.
    Achieves 3-5x faster processing while maintaining identical output.
    """
    from parallel_processor import (
        ParallelDocumentProcessor, 
        build_file_mapping, 
        build_verification_pairs
    )
    import time
    import json
    import io
    
    processor = ParallelDocumentProcessor(max_workers=16)
    
    try:
        logger.info(f"🚀 Background task started for proposal {original_proposal_id} (PARALLEL MODE)")
        verification_results[original_proposal_id] = {}

        extract_dir = proposal_dir / "extracted"
        extract_dir.mkdir(exist_ok=True)

        # ✅ UPDATED: Use filtered extraction (only relevant files)
        pdf_files, kml_files, ignored_count = unzip_files_combined_filtered(
            str(zip_path), 
            str(extract_dir)
        )
        
        if ignored_count > 0:
            logger.info(f"ℹ️  Skipped {ignored_count} non-matching files during processing")

        # ========================================
        # STEP 1: PARALLEL DOCUMENT EXTRACTION
        # ========================================
        file_mapping = build_file_mapping(pdf_files, kml_files)
        
        logger.info(f"📄 Starting parallel extraction of {len([f for f in file_mapping.values() if f])} documents...")
        start_time = time.time()
        
        extracted_docs = processor.extract_documents_parallel(file_mapping)
        
        extraction_time = time.time() - start_time
        logger.info(f"✅ Parallel extraction completed in {extraction_time:.2f}s")
        logger.info(f"📊 Extracted {len(extracted_docs)} document types")

        # ========================================
        # STEP 2: PARALLEL VERIFICATION
        # ========================================
        verification_pairs = build_verification_pairs(extracted_docs)
        
        logger.info(f"🔍 Starting parallel verification of {len(verification_pairs)} pairs...")
        start_time = time.time()
        
        verification_results_data = processor.verify_documents_parallel(
            extracted_docs, 
            verification_pairs
        )
        
        verification_time = time.time() - start_time
        logger.info(f"✅ Parallel verification completed in {verification_time:.2f}s")
        
        # Store verification results
        verification_results[original_proposal_id] = verification_results_data

        # ========================================
        # STEP 3: KML SPECIAL HANDLING
        # ========================================
        if 'kml' in extracted_docs and file_mapping.get('form1'):
            from functions.entity_extraction import extract_form1
            from functions.verification import verify_documents
            
            form1_json = extract_form1(file_mapping['form1'])
            
            if form1_json and 'form1' in form1_json:
                kml_data = extracted_docs['kml']
                form1_data = form1_json['form1']
                geo_coordinates = {
                    "Geographical Co-ordinates (Latitude and Longitude)": form1_data.get("Coordinates")
                }
                verification = verify_documents(kml_data, geo_coordinates, "KML", "FORM1")
                verification_results[original_proposal_id][("KML", "FORM1 Geo Coords")] = verification

        # ========================================
        # STEP 4: SAVE DEBUG DATA
        # ========================================
        extraction_results[original_proposal_id] = extracted_docs
        
        try:
            extracted_data_path = proposal_dir / "extracted_data.json"
            with open(extracted_data_path, "w", encoding="utf-8") as f:
                json.dump(extracted_docs, f, indent=2, ensure_ascii=False)
            logger.info(f"✅ Extracted data saved: {extracted_data_path}")

            verification_data_path = proposal_dir / "verification_results.json"
            serializable_verification = {
                f"{doc1} vs {doc2}": result 
                for (doc1, doc2), result in verification_results[original_proposal_id].items()
            }
            with open(verification_data_path, "w", encoding="utf-8") as f:
                json.dump(serializable_verification, f, indent=2, ensure_ascii=False)
            logger.info(f"✅ Verification results saved: {verification_data_path}")
        except Exception as e:
            logger.error(f"Failed to save debug files: {e}")

        # ========================================
        # STEP 5: OUTPUT GENERATION
        # ========================================
        info_sheet_path = proposal_dir / "Information_Sheet.docx"
        delib_sheet_path = proposal_dir / "Deliberation_Sheet.docx"

        # Deterministic Information Sheet Mapping
        from functions.output_generation12 import build_info_sheet_mapping
        mapping = build_info_sheet_mapping(extracted_docs)
        logger.info(f"Deterministic info sheet mapping generated. Fields filled: {len(mapping)}")

        # Fill Template
        try:
            from functions.output_generation12 import fill_word_with_mapping
            filled_doc_bytes_io, filled_fields_list, missing_list = fill_word_with_mapping(
                str(TEMPLATE_PATH), mapping
            )
            success_info = True
            logger.info(f"✅ Template filled: {len(filled_fields_list)} fields")
            logger.info(f"Missing fields: {len(missing_list)}")
        except Exception as e:
            logger.error(f"Error filling template: {e}", exc_info=True)
            success_info = False
            filled_doc_bytes_io = None
            filled_fields_list = []
            missing_list = []

        if success_info and filled_doc_bytes_io:
            with open(info_sheet_path, "wb") as f:
                f.write(filled_doc_bytes_io.getvalue())
            logger.info(f"✅ Information Sheet saved")
        else:
            logger.error(f"Failed to generate Information Sheet")

        # Determine Missing Files
        found_files = {k: file_mapping.get(k) for k in file_mapping}
        missing_docs = [
            name for name, file_path in {
                "CAF": found_files.get('caf'),
                "Form 1": found_files.get('form1'),
                "NABET": found_files.get('nabet'),
                "Cluster Certificate": found_files.get('cluster'),
                "DMO Site Survey": found_files.get('dmoss'),
                "EMP": found_files.get('emp'),
                "MPA": found_files.get('mpa'),
                "NOC-GP": found_files.get('nocgp'),
                "NOC-Forest": found_files.get('nocforest'),
                "NOC-GSDA": found_files.get('nocgsda'),
                "Kprat": found_files.get('kprat'),
                "GSR": found_files.get('gsr'),
                "Quarry Layout Plan": found_files.get('qlp'),
                "Ownership Document": found_files.get('od'),
                "DSR": found_files.get('dsr'),
                "Affidavit Proposal": found_files.get('western_ghat'),
                "8A": found_files.get('8A'),
                "form1A":file_mapping.get("form1A"),
                "form1B":file_mapping.get("form1B"),
                "pfr":file_mapping.get("pfr"),
                "un_consultant" : file_mapping.get("un_consultant"),
                "un_proponent" : file_mapping.get("un_proponent")
            }.items() if name in REQUIRED_DOC_NAMES and file_path is None
        ]
        missing_files_store[original_proposal_id] = missing_docs
        logger.info(f"Stored {len(missing_docs)} missing files for {original_proposal_id}")

        # Generate Deliberation Sheet
        logger.info("Generating deliberation sheet...")
        serializable_verification = {
            f"{doc1} vs {doc2}": result
            for (doc1, doc2), result in verification_results[original_proposal_id].items()
        }

        try:
            from functions.output_generation12 import gen_delib_sheet
            deliberation_doc_io = gen_delib_sheet(
                extracted_docs, 
                serializable_verification, 
                missing_docs
            )

            if deliberation_doc_io and deliberation_doc_io.getbuffer().nbytes > 0:
                with open(delib_sheet_path, "wb") as f:
                    f.write(deliberation_doc_io.getvalue())
                logger.info(f"✅ Deliberation Sheet saved")
            else:
                logger.error("Failed to generate Deliberation Sheet")
                delib_sheet_path = None

            # Generate MoM
            if delib_sheet_path and delib_sheet_path.exists():
                logger.info("Generating MoM from Info + Deliberation sheets...")
                try:
                    with open(info_sheet_path, "rb") as f:
                        info_io = io.BytesIO(f.read())
                    with open(delib_sheet_path, "rb") as f:
                        delib_io = io.BytesIO(f.read())
                    
                    from functions.output_generation12 import fill_mom_from_info_and_delib
                    mom_io = fill_mom_from_info_and_delib(
                        info_sheet_io=info_io,
                        deliberation_sheet_io=delib_io,
                        mom_template_path="./templates/Stone Quarry - MoM New Format_Draft_051225.docx"
                    )
                    
                    with open(delib_sheet_path, "wb") as f:
                        f.write(mom_io.getvalue())
                    logger.info("✅ MoM generated and saved")
                except Exception as e:
                    logger.error(f"MoM generation failed: {e}", exc_info=True)

        except Exception as e:
            logger.error(f"Deliberation sheet generation failed: {e}", exc_info=True)

        # ========================================
        # STEP 6: AZURE UPLOAD
        # ========================================
        info_sheet_url = None
        delib_sheet_url = None
        try:
            from datetime import datetime, timedelta
            from azure.storage.blob import generate_blob_sas, BlobSasPermissions
            
            if info_sheet_path.exists():
                blob_name_info = f"{original_proposal_id}_info_sheet.docx"
                blob_client_info = blob_service_client.get_blob_client(
                    container=container_name, blob=blob_name_info
                )
                with open(info_sheet_path, "rb") as data:
                    blob_client_info.upload_blob(data, overwrite=True)
                
                sas_info_token = generate_blob_sas(
                    account_name=account_name,
                    container_name=container_name,
                    blob_name=blob_name_info,
                    account_key=account_key,
                    permission=BlobSasPermissions(read=True),
                    expiry=datetime.utcnow() + timedelta(hours=24),
                    response_content_disposition=f'attachment; filename="Information_Sheet_{original_proposal_id.replace("/", "_")}.docx"'
                )
                info_sheet_url = f"https://{account_name}.blob.core.windows.net/{container_name}/{blob_name_info}?{sas_info_token}"
                logger.info("✅ Info sheet uploaded to Azure")

            if delib_sheet_path and delib_sheet_path.exists():
                blob_name_delib = f"{original_proposal_id}_deliberation_sheet.docx"
                blob_client_delib = blob_service_client.get_blob_client(
                    container=container_name, blob=blob_name_delib
                )
                with open(delib_sheet_path, "rb") as data:
                    blob_client_delib.upload_blob(data, overwrite=True)
                
                sas_delib_token = generate_blob_sas(
                    account_name=account_name,
                    container_name=container_name,
                    blob_name=blob_name_delib,
                    account_key=account_key,
                    permission=BlobSasPermissions(read=True),
                    expiry=datetime.utcnow() + timedelta(hours=24),
                    response_content_disposition=f'attachment; filename="Deliberation_Sheet_{original_proposal_id.replace("/", "_")}.docx"'
                )
                delib_sheet_url = f"https://{account_name}.blob.core.windows.net/{container_name}/{blob_name_delib}?{sas_delib_token}"
                logger.info("✅ Deliberation sheet uploaded to Azure")

            blob_urls[original_proposal_id] = {
                "info_sheet_url": info_sheet_url,
                "delib_sheet_url": delib_sheet_url
            }
        except Exception as e:
            logger.error(f"Azure upload failed: {e}")
            blob_urls[original_proposal_id] = {
                "info_sheet_url": None,
                "delib_sheet_url": None,
                "error": str(e)
            }

        # Cleanup ZIP
        safe_id = safe_filename(original_proposal_id)
        zip_path_final = ZIP_UPLOAD_DIR / f"{safe_id}.zip"
        if zip_path_final.exists():
            os.remove(zip_path_final)
        
        logger.info(f"🎉 Background task completed for {original_proposal_id}")

    except Exception as e:
        logger.error(f"❌ Error in background task for {original_proposal_id}: {e}", exc_info=True)
        error_log_path = log_dir / "error.log"
        with open(error_log_path, 'w') as f:
            f.write(f"Error during processing: {e}")
        
        safe_id = safe_filename(original_proposal_id)
        zip_path_final = ZIP_UPLOAD_DIR / f"{safe_id}.zip"
        if zip_path_final.exists():
            os.remove(zip_path_final)
    finally:
        processor.cleanup()
        if lock_file.exists():
            lock_file.unlink(missing_ok=True)

@app.get("/verify-check/{proposal_id:path}")
async def check_verification_status(proposal_id: str):
    """
    Endpoint to check the progress of the verification process.
    """
    decoded_proposal_id = proposal_id
    if decoded_proposal_id in extraction_results:
        if decoded_proposal_id in blob_urls:
            urls = blob_urls[decoded_proposal_id]
            if urls.get("info_sheet_url") and urls.get("delib_sheet_url"):
                status = "Completed"
                return {
                    "proposal_id": decoded_proposal_id,
                    "status": status,
                    "info_sheet_url": urls.get("info_sheet_url"),
                    "delib_sheet_url": urls.get("delib_sheet_url"),
                    "details": "Results available in Azure Blob Storage."
                }
            else:
                status = "Processing"
                return {
                    "proposal_id": decoded_proposal_id,
                    "status": status,
                    "info_sheet_url": None,
                    "delib_sheet_url": None,
                    "details": "Files being generated and uploaded to Azure."
                }
        else:
            return {
                "proposal_id": decoded_proposal_id,
                "status": "Processing",
                "info_sheet_url": None,
                "delib_sheet_url": None,
                "details": "Files being generated."
            }
    else:
        safe_id = safe_filename(decoded_proposal_id)
        proposal_dir = OUTPUT_DIR / safe_id
        if proposal_dir.exists():
            lock_file = proposal_dir / "processing.lock"
            if lock_file.exists():
                status = "Processing"
            else:
                status = "Processing"
            return {
                "proposal_id": decoded_proposal_id,
                "status": status,
                "info_sheet_url": None,
                "delib_sheet_url": None,
                "details": "Files being generated."
            }
        else:
            return {
                "proposal_id": decoded_proposal_id,
                "status": "Not Found",
                "info_sheet_url": None,
                "delib_sheet_url": None,
                "details": "Proposal ID not found or processing not started."
            }

@app.get("/download-file/{proposal_id:path}")
async def download_output_files(proposal_id: str):
    """
    DEPRECATED: Files are now in Azure Blob Storage.
    This endpoint returns the Azure URLs for the files.
    """
    decoded_proposal_id = proposal_id
    if decoded_proposal_id in blob_urls:
        urls = blob_urls[decoded_proposal_id]
        return {
            "info_sheet_url": urls.get("info_sheet_url"),
            "delib_sheet_url": urls.get("delib_sheet_url")
        }
    else:
        raise HTTPException(status_code=404, detail="Output files not found for this proposal ID in Azure.")

@app.get("/download-file/{proposal_id:path}/info")
async def download_info_sheet(proposal_id: str):
    decoded_proposal_id = proposal_id
    if decoded_proposal_id in blob_urls and blob_urls[decoded_proposal_id].get("info_sheet_url"):
        return JSONResponse(content={"url": blob_urls[decoded_proposal_id]["info_sheet_url"]})
    else:
        raise HTTPException(status_code=404, detail="Information Sheet not found in Azure.")

@app.get("/download-file/{proposal_id:path}/delib")
async def download_delib_sheet(proposal_id: str):
    decoded_proposal_id = proposal_id
    if decoded_proposal_id in blob_urls and blob_urls[decoded_proposal_id].get("delib_sheet_url"):
        return JSONResponse(content={"url": blob_urls[decoded_proposal_id]["delib_sheet_url"]})
    else:
        raise HTTPException(status_code=404, detail="Deliberation Sheet not found in Azure.")

@app.get("/extracted-data/{proposal_id:path}")
async def get_extracted_data(proposal_id: str):
    """
    Enhanced endpoint that returns:
    - Extracted data from all documents
    - Missing parameters from Info Sheet
    - Project title, introduction, and facts from Deliberation Sheet
    """
    decoded_proposal_id = proposal_id
    safe_id = safe_filename(decoded_proposal_id)
    proposal_dir = OUTPUT_DIR / safe_id
    
    # Check if extraction results exist
    if decoded_proposal_id not in extraction_results:
        raise HTTPException(status_code=404, detail="No extracted data found for this proposal ID")
    
    # Get existing extracted data
    extracted = extraction_results.get(decoded_proposal_id, {})
    
    # Initialize additional data structures
    missing_parameters = []
    deliberation_data = {}
    missing_files = missing_files_store.get(decoded_proposal_id, [])

    # -----------------------------------------
    # DERIVED AFFIDAVIT BASED FLAGS
    # -----------------------------------------
    unconsulatnt_present = "un_consultant" not in missing_files
    unpropnent_present = "un_proponent" not in missing_files


    affidavit_flags = {
        "Project consultant & EIA Coordinator (Undertaking) – Attached": 
            "Yes" if unconsulatnt_present else "No",

        "Affidavit on EC, EMP, CER Implementation – Attached?": 
            "Yes" if unpropnent_present else "No"
    }

    
    # Extract missing parameters from Information Sheet
    info_sheet_path = proposal_dir / "Information_Sheet.docx"
    if info_sheet_path.exists():
        try:
            missing_parameters = extract_missing_parameters_from_info_sheet(str(info_sheet_path))
            logger.info(f"Extracted {len(missing_parameters)} missing parameters from Info Sheet")
        except Exception as e:
            logger.error(f"Failed to extract missing parameters: {e}")
    
    # Extract data from Deliberation Sheet (MoM)
    delib_sheet_path = proposal_dir / "Deliberation_Sheet.docx"
    if delib_sheet_path.exists():
        try:
            deliberation_data = extract_data_from_deliberation_sheet(str(delib_sheet_path))
            logger.info(f"Extracted deliberation data: Project Title, Introduction, {len(deliberation_data.get('facts_of_proposal', []))} facts")
        except Exception as e:
            logger.error(f"Failed to extract deliberation data: {e}")
    
    return {
        "proposal_id": decoded_proposal_id,
        "extracted_data": extracted,
        # NEW: Missing parameters from Info Sheet
        "missing_parameters": missing_parameters,
        # NEW: Deliberation sheet structured data
        "deliberation_sheet_data": deliberation_data,
        "missing_files": missing_files,
        "missing_files_status": "ALL Documents are Present" if not missing_files else "ALL Documents are not Present",
        "affidavit_checks": affidavit_flags
    }

@app.get("/verification-results/{proposal_id:path}")
async def get_verification_results(proposal_id: str):
    decoded_proposal_id = proposal_id
    if decoded_proposal_id in verification_results:
        # Convert tuple keys to strings for JSON serialization
        serializable_verification = {
            f"{doc1} vs {doc2}": result 
            for (doc1, doc2), result in verification_results[decoded_proposal_id].items()
        }
        return {
            "proposal_id": decoded_proposal_id,
            "verification_results": serializable_verification
        }
    else:
        raise HTTPException(status_code=404, detail="No verification results found for this proposal ID")

@app.get("/all-results/{proposal_id:path}")
async def get_all_results(proposal_id: str):
    decoded_proposal_id = proposal_id
    extracted = extraction_results.get(decoded_proposal_id, {})
    verifications = verification_results.get(decoded_proposal_id, {})
    serializable_verifications = {
        f"{doc1} vs {doc2}": result 
        for (doc1, doc2), result in verifications.items()
    }
    return {
        "proposal_id": decoded_proposal_id,
        "extracted_data": extracted,
        "verification_results": serializable_verifications,
        "info_sheet_url": blob_urls.get(decoded_proposal_id, {}).get("info_sheet_url"),
        "delib_sheet_url": blob_urls.get(decoded_proposal_id, {}).get("delib_sheet_url")
    }

@app.get("/debug-files/{proposal_id:path}")
async def download_debug_files(proposal_id: str):
    decoded_proposal_id = proposal_id
    safe_id = safe_filename(decoded_proposal_id)
    proposal_dir = OUTPUT_DIR / safe_id

    extracted_path = proposal_dir / "extracted_data.json"
    verification_path = proposal_dir / "verification_results.json"

    if not extracted_path.exists() or not verification_path.exists():
        raise HTTPException(status_code=404, detail="Debug files not found. Processing may not be complete.")

    # Create ZIP of debug files
    zip_path = proposal_dir / "debug_files.zip"
    with zipfile.ZipFile(zip_path, "w") as zf:
        zf.write(extracted_path, "extracted_data.json")
        zf.write(verification_path, "verification_results.json")

    return FileResponse(
        path=zip_path,
        filename=f"{safe_id}_debug_files.zip",
        media_type="application/zip"
    )


# ====================================================================
# HELPER FUNCTIONS FOR EXTRACTING DATA FROM GENERATED SHEETS
# ====================================================================

def extract_missing_parameters_from_info_sheet(info_sheet_path: str) -> List[Dict[str, str]]:
    """
    Extract missing parameters from the Information Sheet.
    Returns a list of dictionaries with Sr. No., Parameter, and Remarks.
    """
    doc = Document(info_sheet_path)
    missing_params = []
    
    # Flag to indicate we're in the missing parameters section
    in_missing_section = False
    
    # Look for tables in the document
    for table in doc.tables:
        # Check if this is the missing parameters table
        for row in table.rows:
            cells = row.cells
            if len(cells) >= 3:
                # Check if this is the header row
                cell_text = cells[0].text.strip().lower()
                if "sr. no" in cell_text or "sr.no" in cell_text:
                    in_missing_section = True
                    continue
                
                # If we're in the missing section, extract data
                if in_missing_section:
                    sr_no = cells[0].text.strip()
                    parameter = cells[1].text.strip()
                    remarks = cells[2].text.strip()
                    
                    # Only add if it's actual data (not empty or header)
                    if sr_no and parameter and sr_no.isdigit():
                        missing_params.append({
                            "sr_no": sr_no,
                            "parameter": parameter,
                            "remarks": remarks
                        })
    
    return missing_params

##################################################################################################



def extract_data_from_deliberation_sheet(delib_sheet_path: str) -> Dict[str, Any]:
    """
    Extract Project Title, Introduction, and Facts from Deliberation Sheet (MoM).
    Returns a dictionary with these three keys.
    """
    doc = Document(delib_sheet_path)

    result = {
        "project_title": "",
        "introduction": "",
        "facts_of_proposal": []
    }

    current_section = None
    facts_list = []

    for para in doc.paragraphs:
        text = para.text.strip()

        if not text:
            continue

        lower_text = text.lower()

        # -------------------------------
        # PROJECT TITLE
        # -------------------------------
        if lower_text.startswith("project title:"):
            result["project_title"] = re.sub(r'(?i)project title:\s*', '', text).strip()
            continue

        # -------------------------------
        # INTRODUCTION START
        # -------------------------------
        if "introduction:" in lower_text:
            current_section = "introduction"
            intro_text = re.sub(r'(?i)introduction:\s*', '', text).strip()
            if intro_text:
                result["introduction"] = intro_text
            continue

        # -------------------------------
        # STOP INTRODUCTION WHEN NEW SECTION BEGINS
        # -------------------------------
        if current_section == "introduction" and (
            "project information/details" in lower_text
            or "facts of the proposal" in lower_text
            or "documentary appraisal" in lower_text
            or "deliberations during appraisal" in lower_text
        ):
            current_section = None
            continue

        # -------------------------------
        # COLLECT INTRODUCTION CONTENT
        # -------------------------------
        if current_section == "introduction":
            result["introduction"] += " " + text if result["introduction"] else text
            continue

        # -------------------------------
        # FACTS SECTION START
        # -------------------------------
        if "facts of the proposal" in lower_text and "documentary appraisal" in lower_text:
            current_section = "facts"
            continue

        # -------------------------------
        # STOP FACTS WHEN DELIBERATION STARTS
        # -------------------------------
        if "deliberations during appraisal" in lower_text:
            current_section = None
            continue

        # -------------------------------
        # COLLECT FACTS
        # -------------------------------
        if current_section == "facts":
            match = re.match(r'^(\d+)\.\s+(.+)$', text)

            if match:
                facts_list.append({
                    "item_number": match.group(1),
                    "description": match.group(2)
                })
            elif facts_list:
                # Append continuation lines to last fact
                facts_list[-1]["description"] += " " + text

    # -------------------------------
    # CLEAN INTRODUCTION TEXT
    # -------------------------------
    if result["introduction"]:
        result["introduction"] = re.sub(
            r'(?i)project information/details.*$', 
            '', 
            result["introduction"]
        ).strip()

    result["facts_of_proposal"] = facts_list

    return result


##########################################Not required########################
from fastapi.responses import HTMLResponse
from pathlib import Path

@app.get("/demo", response_class=HTMLResponse)
async def demo_ui():
    html_path = Path("demo.html")
    return html_path.read_text()


################################################################################
# ── SECTION C — ENDPOINTS + BACKGROUND TASK ─────────────────────────────────

@app.post("/seac-process")
async def seac_process(
    proposal_id: str,
    zip_file: UploadFile = File(...),
    background_tasks: BackgroundTasks = BackgroundTasks()
):
    """
    Upload a ZIP containing one or more SEAC PDFs for a given proposal.

    Accepts the same proposal_id format as /validate (slashes allowed —
    pass as a query parameter, e.g. ?proposal_id=SIA/MIN/4353/2024).

    Steps:
      1. Saves the uploaded ZIP to disk.
      2. Kicks off background extraction (Azure DI + LLM per PDF).
      3. Returns immediately — poll /seac-status/{proposal_id} for progress.
    """
    if not zip_file.filename.endswith(".zip"):
        raise HTTPException(status_code=400, detail="File must be a ZIP archive.")

    original_proposal_id = unquote(proposal_id)
    safe_id = safe_filename(original_proposal_id)

    # ── Guard: reject if already processing ──────────────────────────────────
    if seac_processing_status.get(original_proposal_id) == "Processing":
        raise HTTPException(
            status_code=409,
            detail="SEAC processing already in progress for this proposal ID."
        )

    # ── Save ZIP ──────────────────────────────────────────────────────────────
    seac_zip_path = ZIP_UPLOAD_DIR / f"seac_{safe_id}.zip"
    with open(seac_zip_path, "wb") as f:
        f.write(await zip_file.read())
    logger.info(f"[SEAC] ZIP saved: {seac_zip_path}")

    # ── Create proposal output dir ────────────────────────────────────────────
    proposal_dir = OUTPUT_DIR / safe_id
    proposal_dir.mkdir(exist_ok=True)

    # ── Mark as processing and kick off background task ───────────────────────
    seac_processing_status[original_proposal_id] = "Processing"

    background_tasks.add_task(
        _seac_background_task,
        seac_zip_path,
        original_proposal_id,
        proposal_dir,
    )

    return {
        "proposal_id": original_proposal_id,
        "message": "SEAC processing started. Poll /seac-status/{proposal_id} for updates.",
        "status": "Processing",
    }


def _seac_background_task(
    zip_path: Path,
    original_proposal_id: str,
    proposal_dir: Path,
):
    """
    Background task: unzip → extract each SEAC PDF → store results.

    Saves extracted JSON to <proposal_dir>/seac_extracted.json for
    debugging and persistence across server restarts.
    """
    safe_id = safe_filename(original_proposal_id)

    try:
        logger.info(f"[SEAC] Background task started for {original_proposal_id}")

        # ── Step 1: Extract PDFs from ZIP ─────────────────────────────────────
        seac_extract_dir = proposal_dir / "seac_extracted"
        seac_extract_dir.mkdir(exist_ok=True)

        pdf_files, ignored_count = unzip_seac_pdfs(
            str(zip_path),
            str(seac_extract_dir),
        )

        if not pdf_files:
            logger.error(f"[SEAC] No PDFs found in ZIP for {original_proposal_id}")
            seac_processing_status[original_proposal_id] = "Failed"
            seac_results[original_proposal_id] = {
                "error": "No PDF files found inside the uploaded ZIP.",
                "ignored_files_count": ignored_count,
            }
            return

        logger.info(f"[SEAC] {len(pdf_files)} PDF(s) to process, {ignored_count} file(s) skipped")

        # ── Step 2: Extract data from every PDF ───────────────────────────────
        from functions import seac_extraction
        seac_extraction.CURRENT_PROPOSAL_NUMBER = original_proposal_id
        extracted = extract_all_seac_docs(pdf_files)

        # ── Step 3: Store in memory ────────────────────────────────────────────
        seac_results[original_proposal_id] = {
            "documents": extracted,
            "total_files": len(pdf_files),
            "ignored_files_count": ignored_count,
        }
        seac_processing_status[original_proposal_id] = "Completed"
        logger.info(f"[SEAC] Extraction complete for {original_proposal_id}")

        # ── Step 4: Persist to disk (debug / recovery) ────────────────────────
        try:
            seac_json_path = proposal_dir / "seac_extracted.json"
            with open(seac_json_path, "w", encoding="utf-8") as f:
                json.dump(seac_results[original_proposal_id], f, indent=2, ensure_ascii=False)
            logger.info(f"[SEAC] Results saved to disk: {seac_json_path}")
        except Exception as e:
            logger.warning(f"[SEAC] Could not persist results to disk: {e}")

    except Exception as e:
        logger.error(f"[SEAC] Background task failed for {original_proposal_id}: {e}", exc_info=True)
        seac_processing_status[original_proposal_id] = "Failed"
        seac_results[original_proposal_id] = {"error": str(e)}

    finally:
        # ── Cleanup ZIP ────────────────────────────────────────────────────────
        seac_zip = ZIP_UPLOAD_DIR / f"seac_{safe_id}.zip"
        if seac_zip.exists():
            os.remove(seac_zip)
            logger.info(f"[SEAC] ZIP cleaned up: {seac_zip}")


@app.get("/seac-status/{proposal_id:path}")
async def seac_status(proposal_id: str):
    """
    Poll the processing status of a SEAC extraction job.

    Mirrors the /verify-check pattern so the frontend can use
    the same polling loop it uses for the main pipeline.

    Returns:
        status  : "Processing" | "Completed" | "Failed" | "Not Found"
        message : Human-readable description
    """
    decoded_id = proposal_id  # FastAPI already handles path decoding

    status = seac_processing_status.get(decoded_id)

    if status == "Completed":
        return {
            "proposal_id": decoded_id,
            "status": "Completed",
            "message": "SEAC extraction complete. Fetch results from /seac-results/{proposal_id}",
        }
    elif status == "Processing":
        return {
            "proposal_id": decoded_id,
            "status": "Processing",
            "message": "SEAC extraction is in progress. Please check again shortly.",
        }
    elif status == "Failed":
        error_detail = seac_results.get(decoded_id, {}).get("error", "Unknown error")
        return {
            "proposal_id": decoded_id,
            "status": "Failed",
            "message": f"SEAC extraction failed: {error_detail}",
        }
    else:
        # Check if result exists on disk (server restart recovery)
        safe_id = safe_filename(decoded_id)
        seac_json_path = OUTPUT_DIR / safe_id / "seac_extracted.json"
        if seac_json_path.exists():
            return {
                "proposal_id": decoded_id,
                "status": "Completed",
                "message": "Results available (recovered from disk). Fetch from /seac-results/{proposal_id}",
            }

        return {
            "proposal_id": decoded_id,
            "status": "Not Found",
            "message": "No SEAC processing found for this proposal ID. Call /seac-process first.",
        }


@app.get("/seac-results/{proposal_id:path}")
async def seac_get_results(proposal_id: str):
    """
    Return the extracted JSON for all SEAC documents in the uploaded ZIP.

    Response shape:
    {
        "proposal_id": "SIA/MIN/4353/2024",
        "status": "Completed",
        "total_files": 2,
        "ignored_files_count": 1,
        "documents": {
            "SEAC_Meeting_12.pdf": { ...extracted fields... },
            "SEAC_Meeting_13.pdf": { ...extracted fields... }
        }
    }
    """
    decoded_id = proposal_id

    # ── Check in-memory first ─────────────────────────────────────────────────
    if decoded_id in seac_results:
        data = seac_results[decoded_id]
        if "error" in data and "documents" not in data:
            raise HTTPException(
                status_code=500,
                detail=f"SEAC extraction failed: {data['error']}"
            )
        return {
            "proposal_id": decoded_id,
            "status": seac_processing_status.get(decoded_id, "Completed"),
            **data,
        }

    # ── Fall back to disk (server restart recovery) ───────────────────────────
    safe_id = safe_filename(decoded_id)
    seac_json_path = OUTPUT_DIR / safe_id / "seac_extracted.json"
    if seac_json_path.exists():
        try:
            with open(seac_json_path, "r", encoding="utf-8") as f:
                data = json.load(f)
            return {
                "proposal_id": decoded_id,
                "status": "Completed",
                **data,
            }
        except Exception as e:
            logger.error(f"[SEAC] Failed to read persisted results for {decoded_id}: {e}")
            raise HTTPException(status_code=500, detail=f"Failed to read persisted results: {e}")

    # ── Not found ─────────────────────────────────────────────────────────────
    status = seac_processing_status.get(decoded_id, "Not Found")
    if status == "Processing":
        raise HTTPException(
            status_code=202,
            detail="SEAC extraction still in progress. Poll /seac-status/{proposal_id}."
        )

    raise HTTPException(
        status_code=404,
        detail="No SEAC results found for this proposal ID. Call /seac-process first."
    )

################################################################################
################################################################################
if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8504)

this is the code, when I keep the same name of proposal_id which already run once, it wont run again, how to resolve this issue
