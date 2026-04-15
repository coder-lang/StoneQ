from datetime import date
# date = date.today()
import json
import logging
import os
import io
import re
import ast
from typing import Dict, Any, List, Tuple
from datetime import date, datetime
# from functions.fill_infodata import fill_scrutiny_sheet
import pandas as pd
import geopandas as gpd
from docx import Document
from docx.shared import RGBColor
from docx.oxml.ns import qn
from shapely.geometry import (
    Point, LineString, Polygon,
    MultiPoint, MultiLineString,
    MultiPolygon, GeometryCollection
)
from difflib import SequenceMatcher
from Levenshtein import jaro_winkler

from llm_service import extract_with_llm

logger = logging.getLogger(__name__)
logger = logging.getLogger("uvicorn.error")

################################################################################################
# INFO SHEET — DETERMINISTIC CONFIG
################################################################################################




PARAMETER_REFERENCE_PRIORITY = {
    "Project Name": ["caf"],
    "Category": ["form1"],
    "Project Proponent Name": ["form1", "caf"],
    "Project / Activity Cost (in ₹ Lakh)": ["caf"],
    "Permanent Employment Generated": ["caf"],
    "Land use (Agriculture / Industrial / Barren)": ["ss"],

    "Consultant Name": ["nabet"],
    "NABET Certificate No.": ["nabet"],
    "NABET Validity": ["nabet"],

    "Land Gat / Survey No.": ["pfr", "caf"],
    "Village/Town": ["caf", "pfr"],
    "Taluka and District": ["pfr", "caf"],

    "Project Area (Ha.)": ["caf"],
    "Ownership Area as per 7/12 extracts (Ha.)": ["ss","od"],
    "8A Register-Total area (Ha.)": ["ss", "8A"],

    "Nearest Habitat / Village with Distance & Direction": ["ss"],
    "Sensitive Structures in the periphery of 200 m of Stone Quarry (School / Hospital / Residential / National / State Highway / Dam / River / Bridge etc.)": ["ss"],

    "Mining Plan Approval date & Validity": ["mpa"],
    "Minable Total reserve quantity (Brass)": ["mpa"],
    "Excavation quantity (Brass per Year)": ["mpa"],
    "Life of Mine (years)": ["mpa"],
    "Allowed Mining Depth (mts.)": ["mpa"],

    "Total Capital investment (in ₹)": ["emp"],
    "Total Recurring Expenditure (in ₹)": ["emp"],
    "Air & Noise Pollution Control – Total Cost (in ₹)": ["emp"],
    "Water Pollution Control – Total Cost (in ₹)": ["emp"],
    "Environmental Monitoring Cost (Annual) (in ₹)": ["emp"],
    "DMO Site Survey Report (DMO SSR) Letter No. with date":["ss"],
    "DMO report on Cluster Formation (Cluster Certificate / CC) letter no. with date":["cc"],

    "Total CER Budget (in ₹ Lakh)": ["emp"],

    "Morphozone Report from GSDA – Letter No. with Date": ["nocgsda"],
    "Letter from Forest Department – Letter No. with Date": ["nocforest"],
    "Gram Panchayat Tharav and NOC letter – Letter No. with Date": ["nocgp"],
    "Geocoordinates - All corners of Project Site (Latitude & Longitude)": ["form1"],
    "District Survey Report (DSR) Approval Date":["dsr"],
    "Mining Plan Approval date & Validity":["mpa"],
    "Minable Total reserve quantity (Brass)":["mpa"],
    "Excavation quantity (Brass per Year)":["mpa"],
    "Life of Mine (years)":['mpa'],
    "Allowed Mining Depth (mts.)":["mpa"],
    "Air Pollution Control – Total Capital Cost (in ₹)":['emp'],
    "Air Pollution Control – Total Recurring Cost (in ₹)":['emp'],
    "Water Pollution Control – Total Capital Cost (in ₹)":['emp'],
    "Water Pollution Control – Total Recurring Cost (in ₹)":['emp'],
    "Noise Pollution Control – Total Capital Cost (in ₹)":['emp'],
    "Noise Pollution Control – Total Recurring Cost (in ₹)":['emp'],
    "Occupational Health & Safety – Total Cost (in ₹)":["emp"],
    "Solid Waste Management – Total Cost (in ₹)":["emp"],
    "Garland Drainage Execution Cost (in ₹)":["emp"],
    "Proposed Activity (School / Health Centre / Gram Panchayat, etc)": ["emp"],
    "Terrain (Undulating / Plain / Hilly)":["ss"]


}

PARAMETER_JSON_KEYS = {
    "Project Name": ["Name"],
    "Category": ["Category"],
    "Project Proponent Name": ["Project Proponent Name", "Project Proponent Details"],
    "Project / Activity Cost (in ₹ Lakh)": ["Project/Activity Cost"],
    "Permanent Employment Generated": ["Employment"],
    "Consultant Name": ["Consultant Name"],
    "NABET Certificate No.": ["Certificate Number"],
    "NABET Validity": ["Validity Date"],
    "Land Gat / Survey No.": ["land_gat_survey_no","Survey Number"],
    "Village/Town": ["Village/Town", "village_town"],
    "Taluka and District": ["Taluka","Taluka and District"],
    "Project Area (Ha.)": ["Project Area"], 
    "Ownership Area as per 7/12 extracts (Ha.)": ["7/12 Area","Ownership (7/12) Area", "Ownership (7/12) Area in Hectares"],
    "8A Register-Total area (Ha.)": ["8A Register-Total area", "ekun_kshetra"],
    "Nearest Habitat / Village with Distance & Direction": ["Distance from nearest Village/Habitat"],
    "Sensitive Structures in the periphery of 200 m of Stone Quarry (School / Hospital / Residential / National / State Highway / Dam / River / Bridge etc.)": ["Sensitive Structures"],
    "Mining Plan Approval date & Validity": ["Approval Date"],
    "Minable Total reserve quantity (Brass)": ["Minable Total reserve quantity"],
    "Excavation quantity (Brass per Year)": ["Per year excavation quantity"],
    "Life of Mine (years)": ["Life of Mine"],
    "Allowed Mining Depth (mts.)": ["Depth"],
    "DMO Site Survey Report (DMO SSR) Letter No. with date":["Letter No. with Date"],
    "DMO report on Cluster Formation (Cluster Certificate / CC) letter no. with date":["Letter_No_with_Date"],
    "Total Capital investment (in ₹)": ["Total/ Capital investment"],
    "Total Recurring Expenditure (in ₹)": ["Recurring Expenditure Provisions"],
    "Environmental Monitoring Cost (Annual) (in ₹)": ["Environmental Monitoring Recurring Cost"],
    "Total CER Budget (in ₹ Lakh)": ["Total CER budget"],
    "Morphozone Report from GSDA – Letter No. with Date": ["Morphozone Report from GSDA – Letter No. with Date"],
    "Letter from Forest Department – Letter No. with Date": ["Letter from Forest Department – Letter No. with Date"],
    "Gram Panchayat Tharav and NOC letter – Letter No. with Date": ["Gram Panchayat Tharav and NOC letter – Letter No. with Date"],
    "Geocoordinates - All corners of Project Site (Latitude & Longitude)": ["Coordinates"],
    "District Survey Report (DSR) Approval Date":["DSR_Approval"],
    "Mining Plan Approval date & Validity":["Submission to Director of Geology and Mining (Context)"],
    "Minable Total reserve quantity (Brass)":["Minable Total reserve quantity"],
    "Excavation quantity (Brass per Year)":["Per year excavation quantity"],
    "Life of Mine (years)":['Life of Mine'],
    "Allowed Mining Depth (mts.)":["Depth"],
    "Air Pollution Control – Total Capital Cost (in ₹)":['Air Pollution Control Capital Cost'],
    "Air Pollution Control – Total Recurring Cost (in ₹)":['Air Pollution Control Recurring Cost'],
    "Water Pollution Control – Total Capital Cost (in ₹)":["Water Pollution Control Capital Cost"],
    "Water Pollution Control – Total Recurring Cost (in ₹)":["Water Pollution Control Recurring Cost"],
    "Noise Pollution Control – Total Capital Cost (in ₹)":["Noise Pollution Control Capital Cost"],
    "Noise Pollution Control – Total Recurring Cost (in ₹)":["Noise Pollution Control Recurring Cost"],
    "Occupational Health & Safety – Total Cost (in ₹)":["Occupational Health & Safety"],
    "Solid Waste Management – Total Cost (in ₹)":["Solid Waste Management Cost"],
    "Garland Drainage Execution Cost (in ₹)":["Garland Drainage Execution Cost"],
    "Proposed Activity (School / Health Centre / Gram Panchayat, etc)":["Proposed Activity"],
    "Terrain (Undulating / Plain / Hilly)":["Land Type"]



}

SECTION_DISPLAY_NAME = {
    "caf": "CAF",
    "form1": "Form 1",
    "emp": "EMP",
    "nabet": "NABET",

    "mpa": "Mining Plan Approval",
    "ss": "DMO Site Survey Report",
    "od": "Ownership Document – 7/12 extract",
    "nocforest": "Forest Department NOC",
    "nocgsda": "GSDA NOC",
    "nocgp": "Gram Panchayat NOC",
}

from docx.oxml import OxmlElement

def _add_missing_parameters_table(doc, missing_params):
    """
    Safely creates a table under the scrutiny heading
    using python-docx APIs ONLY.
    """
    if not missing_params:
        return

    heading_text = (
        "During the AI scrutiny information on following parameters not available / submitted"
    )

    target_para = None
    for para in doc.paragraphs:
        if heading_text.lower() in para.text.lower():
            target_para = para
            break

    if not target_para:
        return

    # 👉 Create table safely
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"

    # Move table just after heading
    target_para._p.addnext(table._tbl)

    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Sr. No."
    hdr_cells[1].text = "Parameter"
    hdr_cells[2].text = "Remarks"

    # Data rows
    for i, param in enumerate(missing_params, start=1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(i)
        row_cells[1].text = param
        row_cells[2].text = "Information not provided in the attached documents"


def _fetch_value_and_source(data, sections, keys):
    for section in sections:
        block = data.get(section)
        if not isinstance(block, dict):
            continue
        for key in keys:
            if key in block and block[key] not in ("", None, [], {}):
                return block[key], SECTION_DISPLAY_NAME.get(section, "")
    return None, ""


def build_info_sheet_mapping(all_jsons: Dict[str, Any]) -> Dict[str, str]:
    """
    Deterministic Information Sheet mapping.
    NO LLM USED HERE.
    """
    mapping: Dict[str, str] = {}

    for param, sections in PARAMETER_REFERENCE_PRIORITY.items():
        keys = PARAMETER_JSON_KEYS.get(param, [])
        value, source = _fetch_value_and_source(all_jsons, sections, keys)

        if value is not None:
            mapping[param] = f"{value}, {source}" if source else str(value)

    return mapping

def decimal_to_dms(deg, is_lat=True):
    """Convert decimal degrees to DMS string like N18° 27' 06.51"."""
    hemisphere = 'N' if is_lat and deg >= 0 else \
                 'S' if is_lat else \
                 'E' if deg >= 0 else 'W'
    absdeg = abs(deg)
    d = int(absdeg)
    m = int((absdeg - d) * 60)
    s = (absdeg - d - m / 60) * 3600
    return f"{hemisphere}{d}° {m:02d}' {s:05.2f}\""


def extract_latlon_dicts(geom):
    """Convert shapely geometry to list of {'latitude': 'N..', 'longitude': 'E..'}."""
    if geom is None or geom.is_empty:
        return None
    def to_latlon_dict(coords):
        out = []
        for x, y, *rest in coords:
            out.append({
                "latitude": decimal_to_dms(float(y), is_lat=True),
                "longitude": decimal_to_dms(float(x), is_lat=False)
            })
        return out
    if isinstance(geom, Point):
        return to_latlon_dict([geom.coords[0]])
    elif isinstance(geom, LineString):
        return to_latlon_dict(geom.coords)
    elif isinstance(geom, Polygon):
        return to_latlon_dict(geom.exterior.coords)
    elif isinstance(geom, (MultiPoint, MultiLineString, MultiPolygon, GeometryCollection)):
        coords = []
        for g in geom.geoms:
            sub = extract_latlon_dicts(g)
            if sub:
                coords.extend(sub)
        return coords
    return None

def extract_fields_from_word(doc_path: str) -> List[str]:
    """
    Extract all field labels from the 'Parameters' column in Word tables.
    Ignores Sr. No., color, and header cells like 'Parameters', 'Details', etc.
    """
    doc = Document(doc_path)
    fields = set()

    for table in doc.tables:
        # 1) Find the 'Parameters' column index by scanning the first few rows
        param_idx = None
        for row in table.rows[:5]:
            for j, cell in enumerate(row.cells):
                text = cell.text.strip().lower()
                if "parameter" in text:   # matches 'Parameter'/'Parameters'
                    param_idx = j
                    break
            if param_idx is not None:
                break

        # If no Parameters column detected, skip this table
        if param_idx is None:
            continue

        # 2) Collect all non-header items from that column
        for row in table.rows:
            cells = row.cells
            if param_idx >= len(cells):
                continue

            raw_text = cells[param_idx].text.strip()
            if not raw_text:
                continue

            cleaned = re.sub(r"\s+", " ", raw_text).strip(" :-–—\t")

            if not cleaned:
                continue

            # Skip header-like cells
            if cleaned.lower() in ("parameters", "details", "reference document"):
                continue

            fields.add(cleaned)

    return sorted(fields)




# def extract_fields_from_word(doc_path: str) -> List[str]:
#     """Extract field labels from Word template."""
#     doc = Document(doc_path)
#     fields = set()
#     def scan_text(text):
#         matches = re.findall(r"([A-Za-z0-9\-\&\.\,\/\s\(\)]+?):", text)
#         for m in matches:
#             cleaned = m.strip()
#             if 1 < len(cleaned) <= 100:
#                 fields.add(cleaned)
#     for para in doc.paragraphs:
#         scan_text(para.text)
#     for table in doc.tables:
#         for row in table.rows:
#             for cell in row.cells:
#                 scan_text(cell.text)
#     return sorted(fields)


# def llm_map_fields_to_values(fields: List[str], all_jsons: Dict[str, Any]) -> Dict[str, Any]:
#     """Map Word fields to JSON values using LLM."""
#     prompt = f"""
#     You are an expert document integrator.
#     You are given:
#     1️⃣ A list of field labels from a Word template.
#     2️⃣ Structured JSONs extracted from CAF, Form1, NABET, Cluster Certificate, DMO Site Survey, EMP, MPA, NOC-GP, NOC-forest, NOC-GSDA, Kprat, GSR, Quarry Layout Plan (qlp), and Ownership(od) Doc.
#     Task:
#     - Map each Word field to the most accurate value found in these JSONs.
#     - Return a single JSON object.
#     - Preserve units (Ha, ₹, mts, Tons, Years).
#     - If not found, return null.
#     - For each field, include both "value" and "reference" (comma-separated document names).
#     - Output only valid JSON, no text outside JSON.
#     - Do not repeat the value while filling the sheet. eg. In "PP" name repeating, just write once.
#     Special Rules:
#     1. "PP" in the Word field corresponds to all "Project Proponent" names in the JSONs.
#     2. "Gat No." in the Word field corresponds to "Survey Number" in the JSONs.
#     3. "Consultant Name" - EXTRACT FROM NABET document ONLY, not from CAF or Form1. If not available in NABET, return null.
#     4. "DMO report on Cluster Formation" field should be `"Yes"` if a cluster JSON exists or has data; otherwise `"No"`.
#     5. "Total Project Area" field should use the "Area" from Form1, or if missing, from CAF, or if missing, from the "Grand Total Area" in the Cluster Certificate.
#     6. "Excavation Quantity" field should look for values related to mining quantity, output, or production in the EMP or MPA (Mining plan approval) documents.
#     7. "Forest Department Letter" return "Yes" if Forest department letter available in document.
#     8. Ownership (7/12) Area - Look at ownership document (od) and return area 'एकूण क्षेत्र' mentioned in it.
#     9. "Air/ Noise/ Water Pollution mitigation cost" Return all the capital and recurring cost individually.
#     10. "Plantation Cost" Add Capital and Recurring plantation cost and return it.
#     11. Total budget - Total Budget is same as total "Total/ Capital investment"
#     12. NOC's are required if any from (Irrigation /PWD/ Water supply etc.) - Keep it blank.
#     13. Gram Panchayat Resolution and NOC - Return "Yes" if nocgp file is attached in zip file else return "No".
#     14. Forest Department Letter - Return Yes if nocforest file is attached in zip file else return No.
#     15. GSDA-Geological Survey Report - Return Yes if gsr file is attached in zip file else return No.
#     16. Mining plan approval letter - Return Yes if mpa_ file is attached in zip file else return No.
#     17. DMO Site Survey Report - Return 'Yes' if site_survey (ss/dmoss) file available else return 'No'. Mention 'Yes' or 'No' wherever "DMO Site Survey Report" written.
#     18. EC, EMP & CER Implementation (UT): - Always keep it Blank.
#     19. Sensitive Structures in the periphery of 200 Mtrs. of Stone Quarry - Retren all sensitive structures mentioned in dmoss (site survey) file.
#     20. Employment - Look for No. of permanent employment (No.s) in Caf file and return it.
#     21. "Morphozone Report from GSDA – Letter No. with Date" - Return "Yes" from nocgsda document with the required details else return "No".
#     22. "Letter from Forest Department – Letter No. with Date" - Return "Yes" from nocforest document with the required details else return "No".
#     23. "Gram Panchayat Tharav and NOC letter – Letter No. with Date" - Return "Yes" from nocgp docuement with the required details else return "No".
    
#     Return format: {{"field_name": {{"value": "...", "reference": "doc1,doc2"}}, ...}}
    
#     ### FIELD LABELS:
#     {json.dumps(fields, indent=2)}
#     ### DOCUMENT JSONS:
#     {json.dumps(all_jsons, indent=2)}
#     """
#     llm_raw = extract_with_llm(prompt, "")
#     result = _parse_llm_response(llm_raw)
#     return _flatten_value_reference(result)


# def _parse_llm_response(llm_raw):
#     """Parse LLM response to dict."""
#     if isinstance(llm_raw, dict):
#         return llm_raw
#     if isinstance(llm_raw, list):
#         try:
#             return {k: v for d in llm_raw if isinstance(d, dict) for k, v in d.items()}
#         except Exception:
#             return {}
#     if not isinstance(llm_raw, str):
#         llm_raw = str(llm_raw)
#     try:
#         return json.loads(llm_raw)
#     except Exception:
#         match = re.search(r"(\{[\s\S]*\})", llm_raw)
#         if match:
#             json_str = match.group(1)
#             try:
#                 return json.loads(json_str)
#             except Exception:
#                 sanitized = json_str.replace("None", "null").replace("'", '"')
#                 sanitized = re.sub(r",\s*}", "}", sanitized)
#                 sanitized = re.sub(r",\s*]", "]", sanitized)
#                 try:
#                     return json.loads(sanitized)
#                 except Exception:
#                     return {}
#         else:
#             return {}


# def _flatten_value_reference(data: Dict[str, Any]) -> Dict[str, str]:
#     """Flatten {value, reference} to 'value, reference' format."""
#     flattened = {}
#     for field, field_data in data.items():
#         if isinstance(field_data, dict) and "value" in field_data and "reference" in field_data:
#             value = field_data.get("value")
#             reference = field_data.get("reference", "")
#             if value is not None and reference:
#                 flattened[field] = f"{value}, {reference}"
#             elif value is not None:
#                 flattened[field] = str(value)
#             else:
#                 flattened[field] = None
#         else:
#             flattened[field] = field_data
#     return flattened


########################################################################################

DEFAULT_TEMPLATE_PATH = os.getenv(
    "TEMPLATE_PATH",
    "./templates/Stone quarry information sheet_perfect_format.docx"
)


def _norm_label(label: str) -> str:
    """Normalize a label (strip, remove trailing colon, collapse whitespace)."""
    if label is None:
        return ""
    s = str(label).strip()
    s = re.sub(r":\s*$", "", s)
    s = re.sub(r"\s+", " ", s)
    return s


def _find_blanks_by_filled(template_fields: List[str], filled_fields: List[str], mapping: Dict[str, Any]) -> List[str]:
    filled_set = {_norm_label(f) for f in filled_fields}
    blanks: List[str] = []
    for tf in template_fields:
        if _norm_label(tf) not in filled_set:
            val = mapping.get(_norm_label(tf))
            reason = "not filled"
            if val is None or (isinstance(val, str) and not val.strip()):
                reason = "not filled (no value)"
            blanks.append(f"{tf} (reason: {reason})")
    return blanks

#####################################################################################################
#point10 delib
def _resolve_cc_key(cc_dict: dict, *candidates: str):
    """
    Finds a value in cc_dict using flexible key matching.
    Tries: exact → case-insensitive → normalized (no spaces/underscores/hyphens).
    """
    if not isinstance(cc_dict, dict):
        return None

    def normalize(k):
        return re.sub(r"[\s_\-]+", "", k).lower()

    normalized_map = {normalize(k): v for k, v in cc_dict.items()}

    for candidate in candidates:
        # Exact match
        if candidate in cc_dict:
            return cc_dict[candidate]
        # Case-insensitive
        for k, v in cc_dict.items():
            if k.lower() == candidate.lower():
                return v
        # Normalized match (strips spaces, underscores, hyphens)
        if normalize(candidate) in normalized_map:
            return normalized_map[normalize(candidate)]
    return None

# import re
import ast
from Levenshtein import jaro_winkler

# --- Placeholder for Transliteration Logic ---
def normalize_text_for_comparison(text: str) -> str:
    """
    Transliterates Devnagari to Roman script and performs cleanup 
    for robust comparison across scripts, ignoring punctuation.
    """
    if not text:
        return ""
    
    # 1. Transliterate (Placeholder - requires library like unidecode)
    normalized_text = text

    # 2. Lowercase and remove common noise/punctuation
    normalized_text = normalized_text.lower()
    
    # Remove common prefixes/suffixes for names before comparing
    normalized_text = re.sub(r'\b(shri|syed|mohd|mr|late|sri|dr)\.?\s*', '', normalized_text)
    
    # Remove all non-alphanumeric/non-space characters
    normalized_text = re.sub(r'[^a-zA-Z0-9\s]', '', normalized_text) 
    
    # Normalize spaces
    normalized_text = re.sub(r'\s+', ' ', normalized_text).strip() 
    
    return normalized_text

# NOTE: The 'logger', 'all_jsons', and 'extract_with_llm' functions are assumed 
# to be defined and available in your environment.

def _is_real_mismatch(field: str, val_a, val_b) -> bool:
    # --- Initial Checks ---
    a_empty = val_a in [None, "", "N/A"]
    b_empty = val_b in [None, "", "N/A"]
    if a_empty and b_empty:
        return False
    if a_empty or b_empty:
        return True # Real mismatch: one side is missing

    str_a, str_b = str(val_a).strip(), str(val_b).strip()
    field_key = field.lower().strip()

    # 1. Survey Number: normalize by stripping prefixes and suffixes
    if any(kw in field_key for kw in ['survey', 'gat no']):
        def clean_survey_no(s):
            # Strip common prefixes/suffixes (Gut No., Gat No., Part, etc.)
            s = re.sub(r'Gut No\.|Gat No\.|Part', '', s, flags=re.IGNORECASE)
            # Remove all non-alphanumeric characters
            s = re.sub(r'[^a-zA-Z0-9]', '', s).upper()
            return s
            
        norm_a = clean_survey_no(str_a)
        norm_b = clean_survey_no(str_b)
        return norm_a != norm_b

    # 2. Name/Proponent/Owner/Address/Village: Transliteration and Fuzzy Matching
    
    # Identify fields that MUST be resolved by fuzzy matching (TERMINAL)
    is_name_terminal_field = any(kw in field_key for kw in ['name', 'proponent', 'owner'])
    
    # Check if the field is any text field handled here
    if is_name_terminal_field or any(kw in field_key for kw in ['address', 'village']):
        
        norm_a = normalize_text_for_comparison(str_a)
        norm_b = normalize_text_for_comparison(str_b)
        
        # Check 1: Exact match after aggressive cleaning (Script difference fix)
        if norm_a == norm_b:
            logger.info(f"    - Filtered: Script/Punctuation difference (Normalized match)")
            return False 

        # Check 2: Fuzzy Matching (Terminal for Name/Proponent/Owner)
        if is_name_terminal_field:
            FUZZY_THRESHOLD = 0.45 # Recommended high threshold

            if not norm_a or not norm_b:
                 return True # Real mismatch: one side is empty/unparsable

            similarity = jaro_winkler(norm_a, norm_b)
            
            if similarity >= FUZZY_THRESHOLD:
                logger.info(f"    - Filtered: Fuzzy Match (Score: {similarity:.2f})")
                return False # TERMINAL: NOT a real mismatch

            # 🛑 TERMINAL: If a name fails the fuzzy match, it IS a real mismatch.
            logger.info(f"    - REAL MISMATCH: Name failed fuzzy check (Score: {similarity:.2f})")
            return True 
            
        # If the field is 'address' or 'village' and failed the exact match, 
        # it falls through to the LLM (Rule 7) for contextual comparison.
        # This is where the original code would have put its 'fall through' comment.

    # 3. Area fields: parse only numeric part
    if any(kw in field_key for kw in ['project area', 'total area', 'grand total area', 'plot area', 'area']):
        # ... (Area logic remains the same) ...
        def extract_numeric(s):
            s = str(s).replace(',', '')
            match = re.search(r'(\d+(?:\.\d+)?)', s)
            return float(match.group(1)) if match else None
        
        area_a = extract_numeric(str_a)
        area_b = extract_numeric(str_b)
        
        if area_a is not None and area_b is not None:
            return abs(area_a - area_b) > 0.01
        
        return str_a.lower().strip() != str_b.lower().strip()

    # 4. Category: case-insensitive exact match
    if 'category' in field_key:
        return str_a.upper().strip() != str_b.upper().strip()

    # 5. Mineral Type: Check for partial overlap (e.g., 'Basalt' match)
    if any(kw in field_key for kw in ['mineral type', 'extraction type']):
        # ... (Mineral type logic remains the same) ...
        norm_a = normalize_text_for_comparison(str_a)
        norm_b = normalize_text_for_comparison(str_b)
        
        if norm_a in norm_b or norm_b in norm_a:
            logger.info(f"    - Filtered: Mineral type partial match")
            return False
        return True

    # 6. Numeric fields (employment, cost)
    if any(kw in field_key for kw in ['employment', 'cost', 'quantity', 'budget']):
        # ... (Numeric logic remains the same) ...
        cleaned_a = re.sub(r'[^0-9\.]', '', str_a.replace(',', ''))
        cleaned_b = re.sub(r'[^0-9\.]', '', str_b.replace(',', ''))
        
        try:
            val_a_float = float(cleaned_a)
            val_b_float = float(cleaned_b)
            return abs(val_a_float - val_b_float) > 0.01
        except (ValueError, TypeError):
            return cleaned_a != cleaned_b

    # 7. For all other fields (address/village that failed Rule 2, general text), use LLM
    prompt = f"""
    You are an expert document verifier focused on Indian names, addresses, and entities.
    Your task is to determine if these two values for the field '{field}' represent the **SAME ENTITY** or are contextually identical.

    **CRITICAL INSTRUCTIONS:**
    1.  **Ignore Prefixes/Suffixes:** Treat variations like 'Shri.', 'Syed', 'Mohd', 'Mr.', 'Late', etc., as non-substantive.
    2.  **Ignore Script Differences:** Treat names and locations written in English (Roman) and Indian scripts (like Devnagari) as the same if they are clear transliterations.
    3.  **Handle Spelling Variations:** Mark names as the same if the difference is minor (e.g., 'Hameed' vs. 'Hamid', 'Mujmil' vs. 'Muzammil').

    Value A: "{str_a}"
    Value B: "{str_b}"

    Respond ONLY with "Yes" if they are the same entity, and "No" if they are clearly different people or facts.
    """
    try:
        response = extract_with_llm(prompt, "")
        return str(response).strip().lower() != "yes"
    except Exception as e:
        logger.warning(f"LLM check failed for '{field}': {e}. Falling back to strict comparison.")
        return str_a != str_b


# import pandas as pd
from docx.shared import RGBColor
#############################################################################################
#############################################################################################

def gen_delib_sheet(all_jsons: Dict[str, Any], verification_results: Dict[str, Any] = {}, missing_files: List[str] = []) -> io.BytesIO:
    """
    Generates a Deliberation Sheet based on extracted data, verification results, and missing files.
    """
    logger.info("🚀 gen_delib_sheet CALLED")
    mapping = build_info_sheet_mapping(all_jsons)
    # --- Load ESZ villages ---
    ESZ_VILLAGES = set()
    try:
        esz_excel_path = "/home/eytech/ai_mh/API_Stonequarry/data/ESZ_Villages_Maharashtra.xlsx"
        df_esz = pd.read_excel(esz_excel_path, sheet_name=0)
        ESZ_VILLAGES = set(v.strip().lower() for v in df_esz['Village'].dropna().astype(str))
    except Exception as e:
        logger.error(f"Failed to load ESZ Excel: {e}")

    # --- ESZ Village Check ---
    esz_message = "ESZ status: Village not provided or ESZ data unavailable."
    proposed_village = None

    if "form1" in all_jsons and isinstance(all_jsons["form1"], dict):
        form1 = all_jsons["form1"]
        if "Village" in form1:
            proposed_village = str(form1["Village"]).strip()
        else:
            text = " ".join(str(v) for v in form1.values() if isinstance(v, str))
            match = re.search(r'(?i)village\s*[=:,\s]*([a-zA-Z\u0900-\u097F]+)', text)
            if match:
                proposed_village = match.group(1).strip()

    if proposed_village and proposed_village.lower() in ESZ_VILLAGES:
        esz_message = f"The proposed mine village '{proposed_village}' is listed in the Eco-Sensitive Zone (ESZ). Special clearance required."
    elif proposed_village:
        esz_message = f"The proposed mine village '{proposed_village}' is not in the Eco-Sensitive Zone (ESZ) list."

    all_jsons["ESZ_Check"] = esz_message
    logger.info(f"ESZ Check Result: {esz_message}")

    # --- Prepare Missing Files Summary ---
    missing_files_summary_text = ""
    if missing_files:
        missing_list = ", ".join(missing_files)
        missing_files_summary_text = f"The following required documents were NOT FOUND in the submitted proposal: {missing_list}"
    else:
        missing_files_summary_text = "All required documents were found in the submitted proposal."

    # --- Extract template fields ---
    template_path = DEFAULT_TEMPLATE_PATH
    blank_params: List[str] = []
    try:
        template_fields = extract_fields_from_word(template_path)
        logger.info(f"Template fields detected: {len(template_fields)}")
    except Exception as e:
        logger.error(f"Field extraction failed: {e}")
        template_fields = []

    # --- LLM mapping ---
    # try:
    #     mapping = llm_map_fields_to_values(template_fields, all_jsons) if template_fields else {}
    # except Exception as e:
    #     logger.error(f"LLM mapping failed: {e}")
    #     mapping = {}

    # --- Compute blank parameters ---
    # try:
    #     if template_fields and mapping:
    #         _, filled_fields_list = fill_word_with_mapping(template_path, mapping)
    #         blank_params = _find_blanks_by_filled(template_fields, filled_fields_list, mapping)
    #     elif template_fields and not mapping:
    #         blank_params = [f"{tf} (reason: mapping empty)" for tf in template_fields]
    # except Exception as e:
    #     logger.error(f"Fill simulation failed: {e}")
    #     blank_params = []
    
    info_sheet_io = None

    try:
        if template_fields and mapping:
            info_sheet_io = fill_word_with_mapping(template_path, mapping)
            logger.info(f"Info sheet generated: {info_sheet_io is not None}")
        blank_params = []
    except Exception as e:
        logger.error(f"Info sheet generation failed: {e}")
        blank_params = []


    # --- Add NABET validity check ---
    nabet = all_jsons.get("nabet", {})
    if isinstance(nabet, dict) and nabet:
        validity_raw = nabet.get("Validity Date") or nabet.get("Validity Upto")
        if validity_raw:
            from datetime import datetime
            for fmt in ("%B %d, %Y", "%d %B %Y", "%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y"):
                try:
                    expiry_date = datetime.strptime(str(validity_raw).strip(), fmt)
                    break
                except ValueError:
                    continue
            else:
                logger.warning(f"Could not parse NABET validity date: {validity_raw}")
                expiry_date = None
            if expiry_date:
                today = datetime.today()
                is_expired = expiry_date.date() < today.date()
                nabet["is_expired"] = is_expired
                nabet["Expiry Status"] = "Expired" if is_expired else "Valid"
                nabet["Parsed Validity Date"] = expiry_date.strftime("%B %d, %Y")
    
    # >>> PRE-FORMAT LAND HOLDING & KPRAT LINES <<<
# Extract inputs
# >>> PRE-FORMAT LAND HOLDING & KPRAT LINES <<<
    try:
        ss_data = all_jsons.get("ss", {})
        has_ss_land_data = (
            isinstance(ss_data, dict) and
            (
                ss_data.get("7/12 Area") not in (None, "", "Not provided") or
                ss_data.get("8A Register-Total area") not in (None, "", "Not provided")
            )
        )

        if has_ss_land_data:
            ownership_712 = ss_data.get("7/12 Area", "Not provided")
            area_8a = ss_data.get("8A Register-Total area", "Not provided")
        else:
            ownership_712 = all_jsons.get("od", {}).get(
                "Ownership (7/12) Area in Hectares",
                "Not provided"
            )
            area_8a = (
                all_jsons.get("8A", {})
                .get("ekun_kshetra", {})
                .get("value", "Not provided")
            )

        def fmt_area(val):
            """Format an area-like value to 2 decimals where possible, else return as-is."""
            if val in (None, "", "Not provided"):
                return "Not provided"
            try:
                if isinstance(val, str):
                    import re
                    m = re.findall(r"-?\d+(?:\.\d+)?", val)
                    if m:
                        return f"{float(m[0]):.2f}"
                return f"{float(val):.2f}"
            except (ValueError, TypeError):
                return str(val)

        # Build the human-readable line exactly once
        line2_total_land = (
            f"As per 7/12 area is {fmt_area(ownership_712)} Ha. "
            f"and as per 8A area is {fmt_area(area_8a)} Ha."
        )

        # Simple existence check for K-prat key
        kprat_available = "kprat" in all_jsons

        # --- LOGGING (make sure both values are clearly present in logs) ---
        logger.info(f"[DEBUG] kprat key exists: {kprat_available}")
        # Be careful not to dump huge JSON. Limit preview size.
        kprat_preview = str(all_jsons.get('kprat', {}))
        if len(kprat_preview) > 200:
            kprat_preview = kprat_preview[:200] + "...(truncated)"
        # logger.info(f"[DEBUG] kprat content preview: {kprat_preview}")
        logger.info(f"[DEBUG] line2_total_land: {line2_total_land}")

        # (Optional) Avoid print in prod; rely on logger instead.
        # print(f"[PRINT] kprat key exists: {kprat_available}")
        # print(f"[PRINT] line2_total_land: {line2_total_land}")

        # Follow-up line depending on availability
        if kprat_available:
            line3_kprat = "K-Prat is available in the submitted documents."
        else:
            line3_kprat = "The K-prat/TILR map was not provided in the submitted documents."

        logger.info("[MARKER] <<< exiting KPRAT + LAND block")

    except Exception as e:
        logger.exception(f"[ERROR] Failed within KPRAT + LAND block: {e}")
        # Set safe fallbacks to avoid breaking downstream prompt:
        line2_total_land = (
            "The total land holding details from 7/12 and 8A registers were not available."
        )
        line3_kprat = "The K-prat/TILR map was not provided in the submitted documents."
    #New add
    gsda_data = all_jsons.get("nocgsda")
    line7_gsda = (
        "The GSDA NOC has been submitted."
        if gsda_data else
        "The GSDA NOC was not provided in the submitted documents."
    )

    # cc_data = all_jsons.get("cc")
    # line10_cluster = (
    #     "The DMO Cluster Certificate has been submitted."
    #     if cc_data else
    #     "The DMO Cluster Certificate was not provided in the submitted documents."
    # )

    cc_data = all_jsons.get("cc")
    if cc_data and isinstance(cc_data, dict):
        grand_total_area_raw = _resolve_cc_key(
            cc_data,
            "Grand_Total_Area", "grand_total_area", "GrandTotalArea",
            "grand total area", "Grand Total Area"
        ) or ""

        final_remarks = _resolve_cc_key(
            cc_data,
            "Final_Remarks", "final_remarks", "FinalRemarks",
            "final remarks", "Final Remarks"
        ) or ""

        cc_date = _resolve_cc_key(cc_data, "Letter_No_with_Date", "Letter No. with Date") or ""

        area_numeric = None
        if grand_total_area_raw:
            area_match = re.search(r"(\d+(?:\.\d+)?)", str(grand_total_area_raw))
            if area_match:
                area_numeric = float(area_match.group(1))

        if area_numeric is not None:
            area_comparison = (
                f"The Grand Total Area is {grand_total_area_raw}, which is "
                f"{'less than' if area_numeric < 25 else 'not less than'} 25 Ha."
            )
        else:
            area_comparison = (
                f"The Grand Total Area is {grand_total_area_raw}."
                if grand_total_area_raw else
                "The Grand Total Area is not specified."
            )
        date_part = f" dated {cc_date} if cc_date else "

        remarks_part = f" Final Remarks: {final_remarks}" if final_remarks else ""
        line10_cluster = (
            f"The DMO Cluster Certificate {date_part} has been submitted. "
            f"{area_comparison}{remarks_part}"
        )
    else:
        line10_cluster = "The DMO Cluster Certificate was not provided in the submitted documents."

    mpa_data = all_jsons.get("mpa")
    line11_mpa = (
        "The Mining Plan Approval has been submitted."
        if mpa_data else
        "The Mining Plan Approval details were not provided in the submitted documents."
    )

    # --- DSR extraction ---
    dsr_data = all_jsons.get("dsr")
    if dsr_data:
        pp_name = dsr_data.get("Project_Proponent_Name", "Not provided")
        serial_list = dsr_data.get("Corresponding_Serial_Number", [])

        # handle list case
        if isinstance(serial_list, list) and serial_list:
            serial_no = ", ".join(map(str, serial_list))
        else:
            serial_no = "Not provided"

        line4_dsr = (
            f"The DSR Approval page has been submitted. "
            f"The Project Proponent Name '{pp_name}' is listed in the DSR at Serial Number {serial_no}."
        )
    else:
        line4_dsr = "The DSR Approval page and corresponding entry number were not provided in the submitted documents."

    # dmoss_data = all_jsons.get("ss")
    # if dmoss_data:
    #     dmo_date = dmoss_data.get("Letter No. with Date", "Not Provided")
    # else:
    #     dmo_date = "Not Provided"

    # Point 9 — DMO Site Survey Report remarks
    ss_data = all_jsons.get("ss", {})
    if isinstance(ss_data, dict) and ss_data:
        dmo_date = (
            ss_data.get("Letter No. with Date")
            or ss_data.get("Letter_No_with_Date")
            or ss_data.get("letter_no_with_date")
            or ss_data.get("Letter Number with Date")
            or "............"
        )
        
        # Extract remarks — handle list or string
        raw_remarks = (
            ss_data.get("DMO Remarks")
            or ss_data.get("dmo_remarks")
            or ss_data.get("DMO_Remarks")
            or ss_data.get("DMOREMARKS")
            or []
        )
        
        if isinstance(raw_remarks, list):
            remarks_formatted = " ".join(
                f"{i+1}). {r.strip()}" for i, r in enumerate(raw_remarks) if r.strip()
            )
        elif isinstance(raw_remarks, str) and raw_remarks.strip():
            # If already a string, split on common delimiters and number them
            parts = re.split(r'\n|\d+[\)\.]\s*', raw_remarks.strip())
            parts = [p.strip() for p in parts if p.strip()]
            remarks_formatted = " ".join(
                f"{i+1}). {r}" for i, r in enumerate(parts)
            )
        else:
            remarks_formatted = "............"

        line9_dmo = (
            f"As per the submitted DMO Site Survey Report dated {dmo_date}, "
            f"following are the DMO's specific comments/remarks: {remarks_formatted}"
        )
    else:
        line9_dmo = "The DMO Site Survey Report was not provided in the submitted documents."

    
    # Point 18 — Proponent undertaking
    un_proponent_data = all_jsons.get("un_proponent")
    line18_proponent = (
        "The undertaking/affidavit submitted by the Project Proponent regarding implementation of conditions of EC, EMP & CER Plan has been provided."
        if un_proponent_data else
        "The undertaking/affidavit by the Project Proponent regarding implementation of conditions of EC, EMP & CER Plan was not provided in the submitted documents."
    )

    # Point 19 — Consultant undertaking
    un_consultant_data = all_jsons.get("un_consultant")
    line19_consultant = (
        "The undertaking/affidavit submitted by the EIA Coordinator and Project Consultant has been provided."
        if un_consultant_data else
        "The undertaking/affidavit by the EIA Coordinator and Project Consultant was not provided in the submitted documents."
    )


    
    ##################################################################
    # --- Construct the LLM Prompt ---
    deliberation_prompt = f"""
        You are a domain expert preparing an official "Deliberation Sheet" for a stone quarry environmental clearance proposal.
        Your task is to generate the document content based on the input JSONs and missing files.
        Format the output as a Python dictionary string representing the document structure.
        The dictionary should have a single top-level key: "Stone Quarry Deliberation Sheet".
        The value for this key should be another dictionary with keys: "Introduction", "Deliberation and observation of the committee", and "Key Highlights".
        
        - "Introduction": A string containing the introduction text, using values from DOCUMENT JSONS.
          Format: "PP submitted the application for environmental clearance to their proposed stone quarry project having total plot area of <Project_Area> Ha. & maximum per year excavation quantity of mine is <Excavation_Quantity> Brass per year as mentioned at the project details. The project is categorized as <Category> under the EIA Notification, 2006 and its amendments. The project site is located at <Village, Tehsil and District>. The project proponent is <Proponent_Name>."
          If values are missing write sentence in a correct way. 'eg. "The maximum annual excavation quantity for the mine is not specified in the project details"'
        
        - "Deliberation and observation of the committee": A multi-paragraph string containing the Deliberation and observation of the committee text along with the PP submitted details in a SPECIFIC SEQUENCE.
          
          STRUCTURE:
          
          Paragraph 1: "The proposal is appraised on the basis of presentation made & information/documents submitted by PP, environment consultant & concern District Mining Officer."
          
          Paragraph 2: Incorporate the information from 'Missing Documents Summary' if any documents are missing.
          
          Paragraph 3 onwards: Add the PP submitted details:
          
          1. KML File details: For this point only if site survey (ss) data exists, write sensitive structures, road distance, and nearest habitat; else if write "KML file Available".
          2. Total land holding: {line2_total_land}.
          3. K-prat / TILR map: {line3_kprat}.
          4. DSR Approval page & Corresponding entry No details: {line4_dsr}.
          5. Gram Panchayat NOC & Tharav details
          6. Forest NOC details
          7. GSDA NOC details: Using the GSDA NOC content, write a one-sentence English summary covering groundwater condition, morphozone status, and any restrictions or recommendations; base the summary on the following statement: {line7_gsda}
          8. Project Specific any other NOC if required
          9. DMO Site survey report remarks: {line9_dmo}
          10. DMO Certificate on cluster formation details: Based on the following statement, write 2–3 sentences covering: (a) whether the cluster certificate was submitted, (b) the Grand Total Area and whether it is less than 25 Ha, and (c) the final remarks/conclusion on cluster formation and proposal category. Statement: {line10_cluster}
          11. Mining plan details with approval & validity information: Write a one-sentence summary mentioning approval authority, approval date, and validity status based on the following statement: {line11_mpa}
          12. Geological study report (GSR) details
          13. Hydro-geological conditions with respect to ground water levels, aquifer and recharge, slope and stability
          14. Green Belt Development details (budget and number of trees). Also mentioned number of trees planted in 7.5m boundary of proposed quarry.
          15. EMP: Capital Investment & Recurring expenditure details
          16. EMP: Air, Noise & Dust Pollution Mitigation Plan details
          17. CER Implementation Plan, Budget details
          18. Implementation of conditions of EC, EMP & CER Plan: {line18_proponent}
          19. EIA Coordinator and Project Consultant details: {line19_consultant}
          20. Non-Coverage of Proposed Project Land in Western Ghat Ecological Sensitive Area: ESZ statement: Write statement like "The proposed mine village '...' is not included in the List of Villages (Maharashtra) falling under the Western Ghats Eco-Sensitive Zone (ESZ), as per 6th Draft Notification dated 31st July 2024"
        
        IMPORTANT: For items 2 and 3, use the EXACT TEXT shown above. Do NOT rephrase, merge, or alter wording.

        FORMATTING RULES:
        - Each numbered item (1 to 20) MUST be placed on its own line AND must begin with "<number>. ". 
        - You MUST insert a newline before every numbered item. DO NOT place multiple items in the same line under any circumstances
        - Each numbered item MUST be a complete sentence describing what was submitted
        - Start each statement with the document/detail name followed by relevant data
        - If data is not available, write: "The [document name] was not provided in the submitted documents"
        - Do NOT leave sentences incomplete or use standalone colons
        - Use natural, flowing language appropriate for an official document
        - Maintain single spacing between words (no extra spaces)
        - All items under "Deliberation and observation of the committee" must appear as a multi-line text block with each item on a new line. 
          Do NOT compress or merge list items.
        - Do not repeat sentences in introduction.
     
        - "Key Highlights": A list of strings containing key highlights about the project. Include the following information (use "Not provided" if data unavailable):
          1. NABET validity statement: State if NABET certificate is valid or expired with the date.
          2. Mine depth from MPA: Mention the mine depth and highlight if it is less than 25m.
          3. Groundwater depth vs mine depth: State if groundwater depth is less than mine depth.
          4. Morphozone presence: Mention if the project is located in a Morphozone (from NOC-GSDA).
          5. Gram Panchayat NOC validity: State validity date and if it is more than 1 year old.
          6. Forest NOC summary: Mention distance from forest and habitation impact.
          7. K-Prat details: Include survey number, area, and agency information.
          8. Ownership document summary: Include total area and list of landowners.
          9. Quarry Layout Plan area: State the declared area in QLP.
          10. DMO remarks: Provide all remarks points made by the DMO exactly as written in English and do not summarize it. Each remark must be on a new line.
          11. Cluster Certificate grand total: State the grand total area and highlight if it exceeds 25 Ha.
          12. EMP implementation cost: Explicitly state if MPA has no EMP implementation cost.
          13. MPA validity: State the MPA date and compare with current date to mention if validity expired or not.
          14. ESZ statement: Write statement like "The proposed mine village '...' is not included in the List of Villages (Maharashtra) falling under the Western Ghats Eco-Sensitive Zone (ESZ), as per 6th Draft Notification dated 31st July 2024"
          
          Format as: ["1. ...", "2. ...", "3. ..."] etc.
          Keep numbering continuous and use complete sentences.
        
        ---
        ### Output Format (CRITICAL):
        {{
            "Stone Quarry Deliberation Sheet": {{
                "Introduction": "...",
                "Deliberation and observation of the committee": "...",
                "Key Highlights": ["1. ...", "2. ...", "3. ..."]
            }}
        }}
        Do not add any text outside this structure.
        
        ---
        ### Missing Documents Summary (for Deliberation and observation of the committee):
        {missing_files_summary_text}
        
        ---
        ### DOCUMENT JSONS:
        {json.dumps(all_jsons, indent=2)}
        """

    # Call LLM
    llm_response = extract_with_llm(deliberation_prompt, "")
    if not isinstance(llm_response, str):
        llm_response = str(llm_response)
    llm_response = re.sub(r"[\r]+", "", llm_response).strip()
    
    # --- Clean LLM response ---
    llm_response = re.sub(r"```python\s*", "", llm_response)
    llm_response = re.sub(r"```json\s*", "", llm_response)
    llm_response = re.sub(r"```\s*", "", llm_response)
    dict_match = re.search(r"\{.*\}", llm_response, re.DOTALL)
    if dict_match:
        llm_response = dict_match.group(0)

    # --- Create Word Document ---
    doc = Document()
    doc.add_heading("Stone Quarry Deliberation Sheet", level=1)

    parsed_output = None
    sheet_data = {}

    try:
        parsed_output = ast.literal_eval(llm_response)
        if isinstance(parsed_output, dict):
            sheet_data = parsed_output.get("Stone Quarry Deliberation Sheet", {})
    except (SyntaxError, ValueError) as e:
        logger.warning(f"ast.literal_eval failed: {e}. Trying json.loads...")
        try:
            parsed_output = json.loads(llm_response)
            if isinstance(parsed_output, dict):
                sheet_data = parsed_output.get("Stone Quarry Deliberation Sheet", {})
        except json.JSONDecodeError as je:
            logger.error(f"json.loads also failed: {je}")
            parsed_output = None
            sheet_data = {}

    try:
        # Handle Introduction
        if "Introduction" in sheet_data and sheet_data["Introduction"]:
            doc.add_heading("Introduction", level=2)
            doc.add_paragraph(str(sheet_data["Introduction"]))
        
        # Handle Deliberation and observation of the committee with proper formatting
        if "Deliberation and observation of the committee" in sheet_data and sheet_data["Deliberation and observation of the committee"]:
            doc.add_heading("Deliberation and observation of the committee", level=2)
            deliberation_text = str(sheet_data["Deliberation and observation of the committee"])
            
            # Split by newlines (handle both \n and escaped \\n)
            lines = re.split(r'\\n|\n', deliberation_text)
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                # Check if line is a numbered item (starts with digit and dot)
                if re.match(r'^\d+\.\s', line):
                    # Remove the leading number and dot: "1. text" -> "text"
                    text_only = re.sub(r'^\d+\.\s+', '', line)
                    # Remove any duplicate numbering: "1. KML File" -> "KML File"
                    text_only = re.sub(r'^\d+\.\s+', '', text_only)
                    # Add with List Number style (Word will auto-number)
                    doc.add_paragraph(text_only, style='List Number')
                else:
                    # Regular paragraph (not numbered)
                    doc.add_paragraph(line)

        # Key Highlights with orange heading
        if "Key Highlights" in sheet_data and sheet_data["Key Highlights"]:
            heading = doc.add_heading("Key Highlights", level=2)
            for run in heading.runs:
                run.font.color.rgb = RGBColor(255, 165, 0)  # Orange
            kh_items = sheet_data["Key Highlights"]
            for item in (kh_items if isinstance(kh_items, list) else [str(kh_items)]):
                doc.add_paragraph(str(item))
                
    except Exception as e:
        logger.error(f"Failed to process LLM output: {e}")
        doc.add_paragraph("Error: Failed to process LLM output")
        logger.error(f"Sheet data: {sheet_data}")

    # # --- Add Blank Parameters Section ---
    # doc.add_heading("Information Sheet — Blank Parameters", level=2)
    # if blank_params:
    #     doc.add_paragraph("The following Information Sheet parameters were left blank:")
    #     for bp in blank_params:
    #         doc.add_paragraph(bp, style='List Bullet')
    # else:
    #     doc.add_paragraph("No blank parameters detected.")

    # --- Save Document ---
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    logger.info("Deliberation sheet generated successfully")
    return doc_io



###################################################################################################################
###################################################################################################################
def correct_deliberation_format_v2(input_doc_io: io.BytesIO) -> io.BytesIO:
    # import re
    import ast
    doc = Document(input_doc_io)
    full_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
    main_heading = doc.paragraphs[0].text.strip() if doc.paragraphs else "Deliberation Sheet"

    def try_extract_dict(text):
        start = text.find("{")
        end = text.rfind("}")
        if start == -1 or end == -1 or end <= start:
            return None
        candidate = text[start:end + 1]
        try:
            return ast.literal_eval(candidate)
        except:
            try:
                cand2 = re.sub(r"(?<!\\)'", '"', candidate)
                return ast.literal_eval(cand2)
            except:
                return None

    def split_numbered_string(s):
        parts = re.split(r"\n?\s*\d+\.\s+", s)
        return [p.strip() for p in parts if p.strip()]

    parsed_data = try_extract_dict(full_text)
    new_doc = Document()
    new_doc.add_heading(main_heading, level=1)

    if parsed_data:
        for section in ["Introduction", "Deliberations"]:
            text = parsed_data.get(section, "")
            if text:
                new_doc.add_heading(section, level=2)
                new_doc.add_paragraph(str(text).strip())

        items_raw = parsed_data.get("PP submitted the following") or []
        if items_raw:
            new_doc.add_heading("PP submitted the following", level=2)
            item_list = split_numbered_string(items_raw) if isinstance(items_raw, str) else [str(x).strip() for x in items_raw]
            for item in item_list:
                clean_text = re.sub(r"^\d+\.\s*", "", item).strip()
                new_doc.add_paragraph(clean_text, style="List Number")
    else:
        new_doc.add_heading("Formatted Deliberation Sheet (Raw Text)", level=1)
        for line in full_text.splitlines():
            if line.strip():
                new_doc.add_paragraph(line.strip())

    output_doc_io = io.BytesIO()
    new_doc.save(output_doc_io)
    output_doc_io.seek(0)
    return output_doc_io


def value_to_text(val):
    if val is None:
        return ""
    if isinstance(val, (list, tuple)):
        return ", ".join([str(v) for v in val])
    if isinstance(val, dict):
        return json.dumps(val, ensure_ascii=False)
    return str(val)


# def fill_word_with_mapping(doc_path: str, mapping: Dict[str, Any]) -> Tuple[io.BytesIO, List[str]]:
#     doc = Document(doc_path)
#     mapping_clean = {k.strip(): value_to_text(v).strip() for k, v in mapping.items() if v}
#     filled_fields = set()

#     def replace_once_in_text(text: str) -> str:
#         for field, value in mapping_clean.items():
#             if field in filled_fields or not value:
#                 continue
#             pattern = rf"({re.escape(field)}\s*:)"
#             if re.search(pattern, text):
#                 new_text = re.sub(pattern, rf"\1 {value}", text, count=1)
#                 filled_fields.add(field)
#                 return new_text
#         return text

#     for para in doc.paragraphs:
#         new_text = replace_once_in_text(para.text)
#         if new_text != para.text:
#             para.text = new_text

#     for table in doc.tables:
#         for row in table.rows:
#             for cell in row.cells:
#                 new_text = replace_once_in_text(cell.text)
#                 if new_text != cell.text:
#                     cell.text = new_text

#     output = io.BytesIO()
#     doc.save(output)
#     output.seek(0)
#     return output, list(filled_fields)





import io
from typing import Dict, Any, Tuple, List
from docx.oxml.ns import qn

def cell_has_horizontal_merge(cell) -> bool:
    """Detect horizontal merge via w:gridSpan in the cell's XML."""
    tcPr = cell._tc.tcPr
    if tcPr is None:
        return False
    gridSpan = tcPr.find(qn('w:gridSpan'))
    return gridSpan is not None


# ⬇️ Place this helper near other small utilities inside the same file (e.g., above fill_word_with_mapping)
def _expand_reference_docs(ref_text: str) -> str:
    """
    Expand short codes like 'mpa', 'ss', 'emp', 'caf' to full document names
    for the 'Reference Document' column. Unknown codes are kept as-is.
    """
    if not ref_text:
        return ""
    ABBR_MAP = {
        # core ones mentioned
        "mpa": "Mining Plan Approval",
        "ss": "Site Survey Report",
        "dmoss": "DMO Site Survey",
        "emp": "Environmental Management Plan",
        "caf": "Common Application Form",
        # commonly present in your pipeline
        "form1": "Form 1",
        "nocgp": "Gram Panchayat NOC",
        "nocforest": "Forest Department NOC",
        "nocgsda": "GSDA NOC",
        "gsda": "GSDA (Groundwater Study)",
        "gsr": "Geological Study Report",
        "qlp": "Quarry Layout Plan",
        "od": "Ownership Document (7/12)",
        "cluster": "Cluster Certificate",
        "kprat": "K-Prat / TILR Map",
        "nabet": "NABET Certificate",
        "cc": "DMO Cluster Certificate",
    }
    parts = [p.strip() for p in str(ref_text).split(",") if p.strip()]
    expanded = [ABBR_MAP.get(p.lower(), p) for p in parts]
    return ", ".join(expanded)


def fill_word_with_mapping(doc_path: str, mapping: Dict[str, Any]) -> Tuple[io.BytesIO, List[str], List[str]]:
    """
    Safe fill:
    ✔ Split last comma into Details + Reference
    ✔ (Yes/No) auto-set to Yes/No
    ✔ Missing params summary (3 rows)
    ✔ If non-(Yes/No) value missing -> Details = ""
    ✔ SKIP merged header rows to preserve A/B/C/D/E headings
    ✔ 🔴 Keep specific fields BLANK (no LLM data, no "Not found" message)
    """
    doc = Document(doc_path)
    clean_map = {str(k).strip(): v for k, v in mapping.items()}
    filled_fields: List[str] = []
    missing_fields: List[str] = []

    HEADER_TITLES = {
        "Meeting Details",
        "Project Details",
        "Consultant Details",
        "Project Information",
        "Environmental Sensitivity",
        "Project Documents / Approvals",
        "EMP Provisions",
        "CER Provisions",
        "Affidavits, Undertakings, Reports & Letters",
    }

    # 🔴 FIELDS TO KEEP BLANK - regardless of LLM output or missing status
    FIELDS_TO_KEEP_BLANK = {
        "Meeting No. / Item No.",
        "Proposal No.",
        "General / Specific conditions are applicable",
        "Scrutiny fees, if any (in ₹ Lakh)",
    }

    def split_value(val):
        if val is None:
            return None, None
        s = str(val).strip()
        if not s or "," not in s:
            return s if s else None, None
        before, after = s.rsplit(",", 1)
        return (before.strip() or None), (after.strip() or None)

    table = doc.tables[0]

    # ---- Fill Each Parameter Row ----
    for row in table.rows:
        cells = row.cells
        if len(cells) < 2:
            continue

        # Skip section headers (merged across columns)
        if cell_has_horizontal_merge(cells[1]):
            continue

        param_text = cells[1].text.strip()
        # Also skip if matches known header titles
        if param_text in HEADER_TITLES:
            continue

        if len(cells) < 4:
            continue

        param = param_text
        if not param or param == "Parameters":
            continue

        # 🔴 CRITICAL: If field is in KEEP_BLANK list, leave it empty
        if param in FIELDS_TO_KEEP_BLANK:
            cells[2].text = ""
            cells[3].text = ""
            missing_fields.append(param)
            continue

        value = clean_map.get(param)

        if value is not None and str(value).strip():
            if "(Yes/No)" in param:
                cells[2].text = "Yes"
                filled_fields.append(param)
            else:
                details, ref_doc = split_value(value)
                cells[2].text = details if details else " "
                if ref_doc:
                    cells[3].text = _expand_reference_docs(ref_doc)
                filled_fields.append(param)
        else:
            if "(Yes/No)" in param:
                cells[2].text = "No"
            else:
                cells[2].text = " "
                cells[3].text = ""
            missing_fields.append(param)

    # # ---- Missing Parameters Section (Bottom 3 Rows) ----
    # start_index = None
    # for i, row in enumerate(table.rows):
    #     cells = row.cells
    #     if len(cells) >= 2 and cells[1].text.strip().startswith("During the AI scrutiny"):
    #         start_index = i
    #         break

    # if start_index is not None:
    #     list_rows = table.rows[start_index + 1 : start_index + 4]
    #     # 🔴 Exclude fields that are intentionally kept blank from missing summary
    #     missing_clean = [m for m in missing_fields 
    #                     if "(Yes/No)" not in m and m not in FIELDS_TO_KEEP_BLANK]
    #     if len(missing_clean) == 0:
    #         items = ["None", "", ""]
    #     elif len(missing_clean) == 1:
    #         items = [missing_clean[0], "", ""]
    #     elif len(missing_clean) == 2:
    #         items = [missing_clean[0], missing_clean[1], ""]
    #     else:
    #         items = [missing_clean[0], missing_clean[1], ", ".join(missing_clean[2:])]
    #     for row, item in zip(list_rows, items):
    #         if len(row.cells) >= 2:
    #             row.cells[1].text = item

    # Save the filled document first
    temp_io = io.BytesIO()
    doc.save(temp_io)
    temp_io.seek(0)
    final_doc = Document(io.BytesIO(temp_io.getvalue()))


    # ✅ Add missing parameters table in SAME Information Sheet
    missing_clean = [
    m for m in missing_fields
    if "(Yes/No)" not in m and m not in FIELDS_TO_KEEP_BLANK
]

    _add_missing_parameters_table(final_doc, missing_clean)


    # Save final document
    final_io = io.BytesIO()
    final_doc.save(final_io)
    final_io.seek(0)

    return final_io, filled_fields, missing_fields


####################################################################################################

import io
import logging
# import re
from typing import Dict, List, Optional
from docx import Document
from docx.oxml.ns import qn

def fill_mom_from_info_and_delib(
    info_sheet_io: io.BytesIO,
    deliberation_sheet_io: io.BytesIO,
    mom_template_path: str = "./templates/Stone Quarry - MoM New Format_Draft_051225.docx"
) -> io.BytesIO:
    """Build MoM from filled Information Sheet and Deliberation Sheet."""
    
    logger = logging.getLogger(__name__)
    
    def clean_text(text: str) -> str:
        if not text:
            return ""
        return re.sub(r'\s+', ' ', text.strip())
    
    def make_key(text: str) -> str:
        """Normalize text for matching."""
        s = clean_text(text)
        s = re.sub(r'^[A-Z]\.\s+', '', s)
        s = re.sub(r'^\d+\.\s+', '', s)
        return s.lower().strip()
    
    def cell_has_horizontal_merge(cell) -> bool:
        tcPr = cell._tc.tcPr
        if tcPr is None:
            return False
        return tcPr.find(qn('w:gridSpan')) is not None
    
    # ===== EXTRACT INFO SHEET DATA =====
    info_sheet_io.seek(0)
    info_doc = Document(info_sheet_io)
    
    info_data = {}
    info_table = None
    
    # Variables to build Project Title
    project_name = ""
    village = ""
    taluka = ""
    district = ""
    proponent_name = ""
    
    # Find the info sheet table (should have 3+ columns: Sr.No, Parameters, Details)
    for table in info_doc.tables:
        if len(table.rows) > 5 and len(table.rows[0].cells) >= 3:
            info_table = table
            break
    
    if not info_table:
        logger.error("Could not find Info Sheet table")
    else:
        logger.info(f"Found Info Sheet table: {len(info_table.rows)} rows, {len(info_table.rows[0].cells)} columns")
        
        # Extract all parameter->detail mappings (columns 1 and 2)
        for row in info_table.rows:
            if len(row.cells) < 3:
                continue
            
            param_text = clean_text(row.cells[1].text)
            detail_text = clean_text(row.cells[2].text)
            
            if not param_text or not detail_text:
                continue
            
            # Skip headers
            param_lower = param_text.lower()
            if any(h in param_lower for h in ['parameters', 'details', 'sr. no.', 'meeting details', 'project details', 'consultant details', 'project information', 'environmental sensitivity', 'project documents', 'emp provisions', 'cer provisions', 'affidavits']):
                continue
            
            key = make_key(param_text)
            
            # Clean the detail text - remove reference document names
            detail_clean = detail_text
            # Remove common reference suffixes like ", caf", ", form1", etc.
            detail_clean = re.sub(r',\s*(caf|form1|form 1|nabet|cluster|mpa|emp|ss|dmoss|nocgp|nocforest|nocgsda|gsr|qlp|od|kprat|dsr)', '', detail_clean, flags=re.IGNORECASE)
            detail_clean = detail_clean.strip()
            
            info_data[key] = detail_clean
            
            # Capture specific fields for Project Title
            if key == 'project name':
                project_name = detail_clean
            elif key == 'village/town':
                village = detail_clean
            elif key == 'taluka and district':
                # Store the combined taluka and district
                parts = detail_clean.split(',')
                if len(parts) >= 2:
                    taluka = parts[0].strip()
                    district = parts[1].strip()
                else:
                    taluka = detail_clean
                    district = ""
            elif key == 'project proponent name':
                proponent_name = detail_clean
        
        logger.info(f"Extracted {len(info_data)} mappings from Info Sheet")
        
        # Build Project Title in the format:
        # "Project Name (located at) project address (by) Proponent Name"
        project_address = ""
        if village and taluka and district:
            project_address = f"{village}, {taluka}, {district}"
        elif village and taluka:
            project_address = f"{village}, {taluka}"
        elif village:
            project_address = village
        
        project_title = ""
        if project_name and project_name != " ":
            project_title = project_name
            if project_address:
                project_title += f" (located at) {project_address}"
            if proponent_name and proponent_name != " ":
                project_title += f" (by) {proponent_name}"
        else:
            # If no project name, still build with available data
            if project_address:
                project_title = f"(located at) {project_address}"
            if proponent_name and proponent_name != " ":
                project_title += f" (by) {proponent_name}" if project_title else f"(by) {proponent_name}"
        
        if project_title:
            logger.info(f"Built Project Title: {project_title}")
    
    # ===== FIND AND FILL MOM TABLE =====
    mom_doc = Document(mom_template_path)
    
    # Find the MoM "Project Information/Details" table
    # It should have 3 columns and many rows
    mom_table = None
    for i, table in enumerate(mom_doc.tables):
        if len(table.rows) > 20 and len(table.rows[0].cells) >= 3:
            # Check if this table has "Parameters" and "Details" headers
            found_params = False
            found_details = False
            for row in table.rows[:5]:
                for cell in row.cells:
                    text_lower = cell.text.lower()
                    if 'parameter' in text_lower:
                        found_params = True
                    if 'detail' in text_lower:
                        found_details = True
            
            if found_params and found_details:
                mom_table = table
                logger.info(f"Found MoM table at index {i}: {len(table.rows)} rows, {len(table.rows[0].cells)} columns")
                break
    
    if not mom_table:
        logger.error("Could not find MoM Project Information/Details table")
        # Show all tables for debugging
        for i, table in enumerate(mom_doc.tables):
            logger.error(f"Table {i}: {len(table.rows)} rows, {len(table.rows[0].cells) if table.rows else 0} cols")
    else:
        # Fill the MoM table
        filled_count = 0
        for row in mom_table.rows:
            if len(row.cells) < 3:
                continue
            
            # Skip merged header rows
            if cell_has_horizontal_merge(row.cells[1]):
                continue
            
            param_text = clean_text(row.cells[1].text)
            if not param_text:
                continue
            
            # Skip section headers
            param_lower = param_text.lower()
            if any(h in param_lower for h in ['parameters', 'details', 'sr. no.', 'meeting details', 'project details', 'consultant details', 'project information', 'environmental sensitivity', 'project documents', 'emp provisions', 'cer provisions', 'affidavits']):
                continue
            
            key = make_key(param_text)
            if key in info_data:
                row.cells[2].text = info_data[key]
                filled_count += 1
        
        logger.info(f"Filled {filled_count} fields in MoM table")
    
    # ===== FILL PROJECT TITLE =====
    if project_title:
        filled_title = False
        for para in mom_doc.paragraphs:
            para_text = para.text.strip()
            # Match "Project title:" exactly (case-insensitive)
            if para_text.lower().startswith('project title:'):
                para.text = f"Project title: {project_title}"
                logger.info(f"✓ Filled Project Title: {project_title}")
                filled_title = True
                break
        
        if not filled_title:
            logger.warning("Could not find 'Project title:' paragraph to fill")
    
    # ===== EXTRACT DELIBERATION DATA =====
    deliberation_sheet_io.seek(0)
    delib_doc = Document(deliberation_sheet_io)
    
    # Extract Introduction
    intro_text = ""
    for i, para in enumerate(delib_doc.paragraphs):
        if para.text.strip() == "Introduction" and para.style and para.style.name.lower().startswith("heading"):
            # Get next paragraph
            if i + 1 < len(delib_doc.paragraphs):
                intro_text = delib_doc.paragraphs[i + 1].text.strip()
            break
    
    # Fill Introduction in MoM
    if intro_text:
        for para in mom_doc.paragraphs:
            if 'introduction:' in para.text.lower():
                para.text = f"Introduction: {intro_text}"
                break
    
    # Extract Facts (list items under "Deliberation and observation of the committee")
    facts = []
    in_delib_section = False
    for para in delib_doc.paragraphs:
        text = para.text.strip()
        
        if "deliberation and observation of the committee" in text.lower() and para.style and para.style.name.lower().startswith("heading"):
            in_delib_section = True
            continue
        
        if in_delib_section:
            # Stop at next heading
            if para.style and para.style.name.lower().startswith("heading"):
                break
            
            # Check if it's a list item
            style_name = (para.style.name or "").lower() if para.style else ""
            if "list" in style_name and text:
                facts.append(text)
            elif re.match(r'^\d+\.\s+', text):
                facts.append(re.sub(r'^\d+\.\s+', '', text).strip())
    
    # Pad to 20 items
    while len(facts) < 20:
        facts.append("Not provided")
    facts = facts[:20]
    
    # Replace facts in MoM
    facts_heading_idx = None
    for i, para in enumerate(mom_doc.paragraphs):
        if "facts of the proposal" in para.text.lower() and "documentary appraisal" in para.text.lower():
            facts_heading_idx = i
            break
    
    if facts_heading_idx is not None:
        # Find preamble
        preamble_idx = facts_heading_idx + 1
        if preamble_idx < len(mom_doc.paragraphs) and "proposal is appraised" in mom_doc.paragraphs[preamble_idx].text.lower():
            anchor_idx = preamble_idx + 1
        else:
            anchor_idx = facts_heading_idx + 1
        
        # Remove existing numbered items
        i = anchor_idx
        while i < len(mom_doc.paragraphs):
            p = mom_doc.paragraphs[i]
            t = p.text.strip()
            if any(h in t.lower() for h in ["deliberations", "specific conditions", "decision"]):
                break
            if not t or re.match(r'^\d+\.\s+', t) or t.lower() == "not provided":
                p._element.getparent().remove(p._element)
                continue
            break
        
        # Insert new facts in ascending order (1, 2, 3...)
        if anchor_idx < len(mom_doc.paragraphs):
            anchor_para = mom_doc.paragraphs[anchor_idx]
            for idx in range(1, len(facts) + 1):
                anchor_para.insert_paragraph_before(f"{idx}. {facts[idx-1]}")
    
    # Save
    output = io.BytesIO()
    mom_doc.save(output)
    output.seek(0)
    logger.info("✅ MoM generated successfully")
    return output

#####################################################################################################################################
####################################################################################################################################

